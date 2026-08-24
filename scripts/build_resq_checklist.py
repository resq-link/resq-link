import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="D0D7DE", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def make_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def make_first_row_header(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def create_styled_table(doc, headers, data, col_widths, align_cols=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    set_table_borders(table)

    # Style Header Row
    header_row = table.rows[0]
    make_first_row_header(header_row)
    make_row_cant_split(header_row)
    for col_idx, text in enumerate(headers):
        cell = header_row.cells[col_idx]
        cell.text = text
        set_cell_background(cell, "D9E2F3")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = Pt(13)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if align_cols and col_idx in align_cols:
            p.alignment = align_cols[col_idx]
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # Style Data Rows
    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx + 1]
        make_row_cant_split(row)
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(val)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = Pt(13)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if align_cols and col_idx in align_cols:
                p.alignment = align_cols[col_idx]
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x24, 0x29, 0x2F)

    # Set Column Widths
    for row in table.rows:
        for col_idx, width in enumerate(col_widths):
            row.cells[col_idx].width = width

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(8)
    return table

def build_resq_checklist(output_path="RESQ-Link-Testing-Checklist.docx"):
    doc = docx.Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("RESQ-Link End-to-End Testing Checklist")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x2D, 0x59)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(2)
    run_sub = p_sub.add_run("Full system verification — Civilian Mobile, Responder Mobile, Command Center Dispatcher, Super Admin, SMS, Voice, Maps, and Presence")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Date
    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_before = Pt(0)
    p_date.paragraph_format.space_after = Pt(12)
    run_date = p_date.add_run("Prepared August 23, 2026")
    run_date.font.name = "Calibri"
    run_date.font.size = Pt(11)
    run_date.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Instructions Section
    p_inst_heading = doc.add_heading("Instructions for the team", level=2)
    p_inst_heading.paragraph_format.space_before = Pt(10)
    p_inst_heading.paragraph_format.space_after = Pt(4)
    for run in p_inst_heading.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x0F, 0x2D, 0x59)

    instructions = [
        "Work through the phases in order — later phases depend on data and accounts created in earlier ones (the checklist reflects the real-world operational lifecycle of emergency response).",
        "Fill in the Tester column with your initials before executing a test case. Phases can be divided among team members, but Phase 0 (Environment Setup) and Phase 4 (Intake & Triage) should be validated collaboratively.",
        "Mark P (Pass) or F (Fail) in the P/F column. For every Fail, document the defect description, error log, and exact reproduction steps in the Notes column and file a ticket in the project issue tracker.",
        "SMS & Communication Verification: Use ONE dedicated Philippine mobile number (+639...) for live civilian tests, or verify message delivery payloads via the smsLogs collection and dispatch console if SMS gateway test mode is active.",
        "Geolocation & Presence Testing: Perform location tests using physical GPS devices or development simulators configured to Tuguegarao City coordinates (approx. 17.6132° N, 121.7270° E).",
        "Perform testing on a dedicated test Firebase project or clearly isolated staging environment. Several operational flows trigger push notifications, dispatch alerts, and modify active unit availability."
    ]

    for inst in instructions:
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(13)
        run = p.add_run(inst)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2F)

    # Accounts Section
    p_acc_heading = doc.add_heading("Test accounts to prepare (Phases 0–2 creates these)", level=2)
    p_acc_heading.paragraph_format.space_before = Pt(14)
    p_acc_heading.paragraph_format.space_after = Pt(6)
    for run in p_acc_heading.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x0F, 0x2D, 0x59)

    account_headers = ["Role / Account", "Where used & Operational scope", "Created in"]
    account_data = [
        ["Superadmin", "Platform administration, user provisioning, agency configuration, KYC decisions, audit logs (/admin/*)", "Phase 0 (seed script)"],
        ["Command Center Dispatcher", "Intake queue, emergency triage, canonical incident elevation, team dispatch, live map (/command-center/*)", "Phase 0 / Phase 1"],
        ["Field Responder (BFP - Fire)", "Responder mobile app — fire incident alerts, en-route / on-scene status transitions, field assessments", "Phase 1"],
        ["Field Responder (PNP - Police)", "Responder mobile app — law enforcement / traffic dispatch, turn-by-turn navigation, scene reports", "Phase 1"],
        ["Field Responder (CDRRMO / EMS)", "Responder mobile app — medical / rescue dispatch, casualty reporting, post-incident documentation", "Phase 1"],
        ["Civilian A (Mobile + Verified KYC)", "Civilian mobile app — full panic SOS, structured report, live tracking, two-way chat, voice calls", "Phase 2"],
        ["Civilian B (Mobile + Pending/Rejected KYC)", "Civilian mobile app — KYC rejection handling, re-upload workflow, unverified reporting constraints", "Phase 2"],
        ["Civilian C (Phone / Walk-in Caller)", "Dispatcher manual intake — phone/radio emergency logging and offline reporter tracking", "Phase 4"]
    ]
    acc_widths = [Inches(2.2), Inches(3.4), Inches(1.4)]
    create_styled_table(doc, account_headers, account_data, acc_widths)

    # Standard Phase Table Setup
    col_headers = ["#", "Test action", "Expected result", "Tester", "P/F", "Notes"]
    col_widths = [Inches(0.35), Inches(2.35), Inches(2.40), Inches(0.55), Inches(0.40), Inches(0.95)]
    align_center = {0: WD_ALIGN_PARAGRAPH.CENTER, 3: WD_ALIGN_PARAGRAPH.CENTER, 4: WD_ALIGN_PARAGRAPH.CENTER}

    # Helper for adding phase
    def add_phase(title, desc, test_cases):
        p_h1 = doc.add_heading(title, level=1)
        p_h1.paragraph_format.space_before = Pt(14)
        p_h1.paragraph_format.space_after = Pt(3)
        for run in p_h1.runs:
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x0F, 0x2D, 0x59)

        if desc:
            p_d = doc.add_paragraph()
            p_d.paragraph_format.space_before = Pt(0)
            p_d.paragraph_format.space_after = Pt(6)
            run_d = p_d.add_run(desc)
            run_d.font.name = "Calibri"
            run_d.font.size = Pt(9.5)
            run_d.font.italic = True
            run_d.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        table_data = []
        for idx, (action, expected) in enumerate(test_cases, 1):
            table_data.append([str(idx), action, expected, "", "", ""])

        create_styled_table(doc, col_headers, table_data, col_widths, align_center)

    # Phase 0
    add_phase(
        "Phase 0 — Environment Setup, Firebase Configuration, and Seeding",
        "Execute once on the test environment prior to functional validation. Verifies environment variables, rule deployments, seed scripts, and core client startup.",
        [
            ("Copy .env.example files to .env.local in apps/resq-link-web-app and .env in apps/civilian-mobile-app and apps/responder-mobile-app; populate Firebase client credentials, Admin SDK private key, Agora App ID/Certificate, and Mapbox tokens.",
             "All applications boot cleanly with zero missing environment variable or configuration runtime warnings."),
            ("Execute npm install at the repository root to link npm workspaces across all packages, web apps, mobile apps, and Cloud Functions.",
             "Workspaces dependency graph resolves cleanly; native patch-package scripts apply with 0 errors."),
            ("Deploy Cloud Firestore security rules, Firebase Storage rules, and Realtime Database presence rules (firebase deploy --only firestore:rules,storage,database).",
             "Firebase CLI returns deployment success; security rules are compiled and active on the test database instance."),
            ("Deploy Firestore composite query indexes (firebase deploy --only firestore:indexes).",
             "All required composite indexes for emergencies, incidents, teams, and audit collections build and report enabled."),
            ("Seed the initial Super Admin account: npx ts-node packages/firebase/scripts/create-first-admin.ts.",
             "Super Admin document created in Firestore admins collection with role 'super_admin'; credentials validated."),
            ("Seed default incident categories and dispatch routing rules: npx ts-node packages/firebase/scripts/seed-incident-type-rules.ts.",
             "Firestore incidentTypeRules collection populated with Fire, Medical, Crime, Vehicular Accident, Rescue, and Disaster rules."),
            ("Start the consolidated web portal (npm run dev --workspace=apps/resq-link-web-app) and open http://localhost:3000.",
             "RESQ-Link landing portal loads with login entry points for both Command Center and Super Admin workspaces."),
            ("Start the Civilian Mobile App and Responder Mobile App in Expo dev mode (npm start).",
             "Both mobile apps boot up and display their respective authentication / landing screens on device or emulator."),
            ("Verify external service configurations: Agora Voice RTC endpoint (/api/agora/token), Mapbox/Leaflet tiles, and SMS gateway logger.",
             "Agora token generator returns valid JSON response; web and mobile map components load base map tiles.")
        ]
    )

    # Phase 1
    add_phase(
        "Phase 1 — Super Admin Management and System Governance (/admin/*)",
        "Verify administrative provisioning, agency setup, dispatcher/responder account lifecycle, civilian registry, and immutable audit logs.",
        [
            ("Log in as Super Admin on the web portal (/login).",
             "Authentication succeeds; system validates role against admins/{uid} and routes to Super Admin Dashboard (/admin/dashboard)."),
            ("Verify Dashboard KPI cards (Total Civilians, Registered Responders, Active Dispatchers, Pending KYC, Total Incidents).",
             "Metric counters display accurate live figures matching existing Firestore database records."),
            ("Agency Management (/admin/agencies): Create a new emergency response agency (e.g., BFP Tuguegarao, PNP Station 1, CDRRMO Rescue 116).",
             "Agency record saves in Firestore agencies collection; agency becomes immediately selectable in team creation."),
            ("Edit Agency details (hotline number, primary contact, active jurisdiction).",
             "Modifications persist upon page refresh and correctly update agency metadata across dispatcher views."),
            ("Dispatcher Management (/admin/dispatchers): Provision a new Command Center Operator account.",
             "Auth account and dispatchers / commandCenters documents created; dispatcher appears in active operators table."),
            ("Responder Provisioning (/admin/responders): Provision field responder accounts assigned to specific agencies (BFP, PNP, CDRRMO).",
             "Responder accounts created with designated agency role, call sign, unit ID, and operational status."),
            ("Civilian Registry (/admin/civilians): Search and filter civilian users by verification status, name, or phone number.",
             "Civilian records list correctly with KYC status badges (Verified, Pending, Unverified); details modal displays profile."),
            ("Account Deactivation: Deactivate a dispatcher or responder account, attempt login, then reactivate.",
             "Login is blocked with an 'Account Deactivated' notice while disabled; login succeeds immediately upon reactivation."),
            ("Audit Logs (/admin/audit): Review system audit trail for recent logins, account creations, role updates, and deletions.",
             "Audit entries display timestamp, actor UID, action type, IP address, and target record ID."),
            ("Theme & Display Preferences: Toggle Admin theme between Light and Dark mode using the AdminThemePicker.",
             "UI theme dynamically updates stylesheet variables and persists selection in browser local storage.")
        ]
    )

    # Phase 2
    add_phase(
        "Phase 2 — Civilian Identity, Registration, KYC Verification, and Privacy Consent",
        "Validate civilian mobile registration, mandatory legal privacy consents, document upload, Super Admin KYC review, and password recovery.",
        [
            ("On Civilian Mobile App, open Registration and submit new user details (Full Name, Email, PH Mobile +639..., Password, Address, Emergency Contact).",
             "Input fields validate formatting; user account created in Firebase Auth and Firestore users collection."),
            ("Verify Mandatory Legal Consent: Attempt registration without ticking Terms of Use, Privacy Policy, and Data Privacy Notice checkboxes.",
             "Registration button remains disabled; tapping legal links opens the full markdown legal policy viewer modals."),
            ("Attempt duplicate registration using an existing email address or registered mobile number.",
             "Form displays inline validation error ('Email or phone number already in use'); duplicate record creation prevented."),
            ("Log in as newly registered civilian on mobile app; close and reopen app.",
             "Login succeeds; user session persists across app restarts without requiring re-authentication."),
            ("Submit KYC Verification Documents: From civilian profile, capture and upload Government ID photo and live selfie.",
             "Images upload securely to Firebase Storage (kyc/{uid}/*); user record updates to 'pending_kyc' status."),
            ("Super Admin KYC Queue (/admin/kyc): Locate the submitted civilian verification request.",
             "KYC review drawer displays applicant info, high-resolution ID photo, selfie, and submission timestamp."),
            ("Reject KYC Submission: Reject a test civilian KYC with a documented reason (e.g. 'Blurry ID document').",
             "Status updates to 'rejected'; Civilian Mobile App immediately displays rejection banner and allows document re-upload."),
            ("Approve KYC Submission: Super Admin approves valid civilian KYC submission.",
             "User status updates to 'verified'; Civilian Mobile App displays verified checkmark badge on profile and home screen."),
            ("Test Password Recovery: Submit 'Forgot Password' request with valid registered email from civilian login screen.",
             "Firebase password reset email received; reset link allows updating password; new password functions on login.")
        ]
    )

    # Phase 3
    add_phase(
        "Phase 3 — Emergency Reporting and Civilian Mobile Experience",
        "Test panic triggers, structured emergency reporting forms, GPS geocoding, photo evidence uploads, offline queueing, and live tracking.",
        [
            ("Trigger Panic SOS / Shake-to-Emergency: Shake the device or press the SOS panic button on Civilian home screen.",
             "Emergency confirmation modal appears with 5-second audible/vibration countdown and one-tap Cancel option."),
            ("Let SOS timer complete without cancelling.",
             "Instant panic report generated with auto-captured GPS coordinates and submitted to Command Center intake."),
            ("Open Structured Emergency Form: Select emergency category (Fire, Medical, Crime, Vehicular Accident, Disaster, Rescue).",
             "Form expands with category-specific fields, severity questions, and required information indicators."),
            ("GPS Location Capture & Map Pin Adjustment: Verify auto-detected location and manually drag pin to precise landmark.",
             "Latitude, longitude, detected barangay, and nearest landmark populate accurately in report payload."),
            ("Attach Photo Evidence: Capture 1–3 photos using mobile camera or select from image gallery.",
             "Image thumbnails preview with removal buttons; images compress and attach to report data structure."),
            ("Specify Casualties & Narrative: Enter estimated number of injured/trapped individuals and emergency description.",
             "Form validates required fields and enables the 'Submit Emergency Report' action button."),
            ("Submit Emergency Report & Verify Confirmation.",
             "Report created in emergencies collection; app navigates to Live Incident Tracking screen with unique tracking ID."),
            ("Offline Resilience Test: Enable Airplane Mode on mobile, submit emergency report, then disable Airplane Mode.",
             "App warns user of offline state, queues the emergency locally, and auto-submits once network connection is restored."),
            ("Civilian Emergency History: Open History tab (/history) to review past submitted emergencies.",
             "List displays previous reports with timestamps, incident categories, response summaries, and final resolution statuses.")
        ]
    )

    # Phase 4
    add_phase(
        "Phase 4 — Command Center Intake, Triage, and Deduplication (/command-center/*)",
        "Validate incoming emergency alerts, dispatcher manual intake, priority scoring, duplicate detection, canonical incident elevation, and rejection.",
        [
            ("Log in as Command Center Dispatcher (/login) and navigate to Emergency Intake (/command-center/intake).",
             "Intake dashboard loads showing active incoming queue, priority filters, and sound notification controls."),
            ("Trigger an emergency from Civilian Mobile App while Dispatcher Intake is open.",
             "Audible dispatch alert sounds; new emergency record appears at top of intake queue highlighted with 'NEW' badge."),
            ("Manual Intake (/command-center/intake/new): Log an emergency received via Phone Call, Radio, or Walk-in.",
             "Dispatcher completes intake form (caller name, contact, incident type, location pin, narrative); saved with source 'phone_call'."),
            ("Incident Triage: Dispatcher inspects report details and assigns operational priority (Critical, High, Medium, Low).",
             "Priority badge updates immediately; card color reflects severity (Red = Critical, Orange = High, Yellow = Medium)."),
            ("Duplicate Detection & Grouping: Submit two reports for the same incident category within 500 meters of each other.",
             "System flags potential duplicate report; dispatcher can link multiple citizen reports under one parent incident."),
            ("Elevate Report to Canonical Incident: Dispatcher clicks 'Create / Elevate Incident' from the validated report.",
             "Canonical incidents document created; links source emergency ID; assigns official tracking number (INC-2026-XXXX)."),
            ("False Alarm / Invalid Report Rejection: Dispatcher rejects a test spam report with a mandatory reason.",
             "Report status transitions to 'rejected'; civilian mobile tracking reflects rejection explanation; report removed from active queue."),
            ("Verify Emergency Intake Badge Count on Navigation Sidebar.",
             "Intake badge counter matches total unacknowledged incoming reports; decreases as reports are elevated or rejected.")
        ]
    )

    # Phase 5
    add_phase(
        "Phase 5 — Incident Dispatch, Agency Routing, and Resource Management",
        "Test automated agency recommendations, team dispatch, individual responder assignment, vehicle allocation, and dynamic reassignment.",
        [
            ("Open newly elevated incident in Incident Management (/command-center/incidents/[id]).",
             "Incident detail view displays full narrative, photos, caller info, location map, and Agency Recommendation panel."),
            ("Verify Agency Auto-Recommendation Matrix: Check recommended agencies for Fire (BFP + EMS) vs Vehicular Accident (PNP + Ambulance).",
             "Suggested agency tags automatically match predefined incident type routing rules seeded in Phase 0."),
            ("Operational Team Assignment: Assign an available response team (e.g. 'BFP Alpha Unit 1') to the incident.",
             "Team status updates from 'Available' to 'Assigned'; all team members linked to incident dispatch roster."),
            ("Individual Responder Assignment: Assign specific individual responders by call sign or agency badge.",
             "Responders added to assignedResponders list; dispatch push alerts triggered to their mobile devices."),
            ("Vehicle & Equipment Allocation (/command-center/resources): Allocate emergency vehicles (Fire Engine 1, Ambulance 2) to the incident.",
             "Vehicle operational status switches to 'Deployed'; equipment list links to active incident record."),
            ("Multi-Agency Escalation: Add PNP Police escort and additional CDRRMO rescue unit mid-operation.",
             "Incident reflects multi-agency deployment; secondary agency dispatchers and responders receive assignment alerts."),
            ("Live Team Reassignment: Reassign incident from Team 1 to Team 2 due to unit unavailability.",
             "Team 1 released back to 'Available'; Team 2 alerted; reassignment reason logged in incident chronological history."),
            ("Verify Dispatch Status Synchronized to Civilian Tracking View.",
             "Civilian mobile app timeline updates to 'Responders Dispatched' with names of assigned response agencies.")
        ]
    )

    # Phase 6
    add_phase(
        "Phase 6 — Field Responder Operations and Mobile Lifecycle",
        "Validate responder mobile shift toggle, dispatch alert receipt, accept/decline actions, turn-by-turn navigation, status updates, and PIR completion.",
        [
            ("Log in to Responder Mobile App with assigned agency credentials (e.g. BFP Responder account).",
             "Responder lands on Responder Dashboard displaying shift status, agency emblem, and assigned call sign."),
            ("Shift & Availability Toggle: Switch status between 'On Duty', 'Available', and 'Off Duty'.",
             "Status updates in Firestore and Realtime Database; reflects in real time on Command Center dispatcher console."),
            ("Receive Dispatch Push Alert: Dispatch an active incident to the responder from Command Center.",
             "Responder app triggers high-priority alert sound/vibration; full-screen assignment popup appears with incident preview."),
            ("Accept Dispatch Assignment: Responder taps 'Accept Assignment'.",
             "Incident assignment status becomes 'accepted'; dispatcher console indicates responder acknowledged; navigation map opens."),
            ("Decline Assignment Flow: Dispatch another incident and tap 'Decline', selecting reason 'Vehicle Breakdown'.",
             "Assignment cleared from responder; Command Center alerted immediately to re-dispatch; decline reason logged in history."),
            ("Incident Details & Field Navigation: View civilian notes, hazard photos, and tap 'Navigate to Scene'.",
             "Map opens with GPS route navigation to the emergency coordinates in Tuguegarao City."),
            ("Status Transition — 'Mark En Route': Responder taps 'En Route' button.",
             "Status updates to 'en_route'; timestamp recorded; civilian app updates to 'Responders are on the way'."),
            ("Status Transition — 'Mark On Scene': Responder taps 'Arrived On Scene'.",
             "Status updates to 'on_scene'; timestamp recorded; Command Center map marker icon turns blue/on-scene."),
            ("Field Assessment & Photo Upload: Upload on-scene damage photo and enter tactical notes.",
             "Photos and field notes synchronize in real time to Command Center incident management view."),
            ("Post-Incident Report (PIR) & Case Closure: Fill in casualties treated, resources used, narrative summary, and tap 'Complete Incident'.",
             "Incident marked 'resolved'; PIR stored; responder and allocated vehicles automatically reset to 'Available'.")
        ]
    )

    # Phase 7
    add_phase(
        "Phase 7 — Real-Time Maps, Geolocation, and Responder Presence",
        "Test Command Center Leaflet/Mapbox map layers, live GPS streaming via Realtime Database, civilian tracking map, and layer filtering.",
        [
            ("Open Command Center Operations Map (/command-center/map).",
             "Map renders centered on Tuguegarao City; displays active incident markers, agency headquarters, and responder units."),
            ("Incident Marker Verification: Click on active incident markers on the map.",
             "Popup displays Incident ID, category icon, priority color, elapsed response time, and assigned units."),
            ("Real-Time GPS Presence: With Responder app active in foreground, simulate moving responder coordinates.",
             "Responder vehicle marker moves smoothly on Command Center map in near real time via Firebase Realtime Database."),
            ("Civilian Live Unit Tracking: On Civilian Mobile App live tracking screen, observe assigned unit approaching.",
             "Civilian sees responder icon on map with live distance (meters) and estimated time of arrival (ETA)."),
            ("Map Layer & Agency Filtering: Filter map display by Agency (BFP only, PNP only, CDRRMO only) and Priority (Critical only).",
             "Map dynamically filters visible markers without full-page reloads or tile flickering."),
            ("Search Barangay / Landmark: Use the map location search bar to locate specific Tuguegarao barangays (e.g. Carig Sur, Centro).",
             "Map smoothly pans and zooms to the selected barangay coordinates."),
            ("Map Fallback Behavior: Simulate offline/blocked Mapbox tile server.",
             "Map falls back cleanly to OpenStreetMap tile provider with no unhandled script exceptions."),
            ("Responder Presence Heartbeat Timeout: Close Responder Mobile App completely.",
             "Presence system marks responder offline after disconnect timeout; map marker updates to offline state.")
        ]
    )

    # Phase 8
    add_phase(
        "Phase 8 — Incident Communication: In-App Messaging and Agora Voice Calls",
        "Verify real-time two-way civilian-dispatcher chat, responder coordination threads, Agora voice token generation, and audio call controls.",
        [
            ("Civilian-to-Dispatcher Chat: From Civilian Mobile App active incident screen, send a chat message to Command Center.",
             "Message appears in real time in Command Center incident chat drawer with timestamp and sender tag."),
            ("Dispatcher Reply: Dispatcher sends a reply message with an advisory instruction.",
             "Civilian receives message immediately in chat thread with unread indicator and sound chime."),
            ("Responder Field Chat: Send messages and photos between Responder Mobile App and Command Center.",
             "Dedicated operational channel displays field coordination messages with high-resolution image previews."),
            ("Agora Voice Token Generation (/api/agora/token): Dispatcher clicks 'Call Civilian' from browser console.",
             "Next.js API route generates dynamic Agora RTC channel token; logs active session in callSessions collection."),
            ("Voice Call Connection (Web to Mobile): Civilian Mobile App receives incoming call overlay; tap 'Accept Call'.",
             "Two-way crystal-clear audio stream establishes; voice transmits cleanly between web browser and mobile device."),
            ("Audio Call Controls: Test Mute/Unmute microphone, Speakerphone toggle, and End Call on both devices.",
             "Microphone mute state reflects accurately; audio halts when muted; call terminates cleanly on both ends upon hanging up."),
            ("Call History & Duration Logging: Review incident communication log after call completion.",
             "Call session record displays start time, end time, total call duration, participant UIDs, and termination status.")
        ]
    )

    # Phase 9
    add_phase(
        "Phase 9 — SMS Integration, Public Advisories, and Community Broadcasts",
        "Validate automated SMS alerts, dispatcher manual SMS tool, public weather/disaster advisory creation, mobile banner broadcasts, and expiry.",
        [
            ("Automated Dispatch SMS: Dispatch an incident for a civilian with a valid Philippine mobile number (+639...).",
             "SMS gateway (or smsLogs in test mode) records outgoing SMS: 'RESQ-Link: Responders have been dispatched to your location...'."),
            ("Automated Resolution SMS: Mark the incident as resolved in Command Center.",
             "Resolution SMS sent to civilian stating the emergency response has been concluded."),
            ("Manual Dispatcher SMS (/command-center/sms): Send direct SMS notification to a specific civilian contact number.",
             "Message logged in SMS outgoing log with recipient number, message body, timestamp, and delivery status."),
            ("SMS Error Handling: Send SMS to an invalid/unreachable phone number.",
             "System logs graceful failure in smsLogs without throwing unhandled exceptions or disrupting dispatch UI."),
            ("Create Public Advisory (/command-center/advisories): Create city-wide advisory (e.g. 'Typhoon Signal No. 2 Alert — Flood Warning').",
             "Form validates title, category, severity (Critical, Warning, Advisory), affected barangays, and auto-expiry date."),
            ("Advisory Broadcast & Push Notification: Confirm broadcast in the Broadcast Confirmation Modal.",
             "Advisory saved to advisories collection; push notification broadcast triggered to all registered civilian devices."),
            ("Civilian Mobile Advisory Feed: Open Civilian Mobile App home screen and Advisories tab (/advisories).",
             "Emergency advisory banner displays prominently on home screen; full advisory details viewable with actionable safety tips."),
            ("Advisory Expiry / Deactivation: Dispatcher expires or revokes an active advisory.",
             "Advisory disappears from civilian home banner and moves to archived advisory history.")
        ]
    )

    # Phase 10
    add_phase(
        "Phase 10 — Footage Requests, Incident History, and Operational Analytics",
        "Test traffic CCTV footage requests, advanced incident history search/filters, chronological audit reviews, and analytics charts.",
        [
            ("Create CCTV Footage Request (/command-center/footage-requests): Submit request for traffic/CCTV footage.",
             "Request saved to footageRequests collection with camera location, intersection, time window, and requesting officer details."),
            ("Update Footage Request Status: Dispatcher marks footage request as 'Footage Available' and attaches archive link.",
             "Status updates in real time; requesting agency receives update notification."),
            ("Incident History Search (/command-center/history): Filter closed incidents by Date Range, Barangay, Agency, and Priority.",
             "Data table filters instantaneously; search query locates specific incidents by tracking number or keyword."),
            ("Incident Chronology Audit: Open historical incident record and inspect chronological event log.",
             "Complete audit trail viewable: Intake timestamp -> Triage -> Dispatch -> En Route -> On Scene -> PIR -> Closure."),
            ("Operational Analytics Overview (/command-center/overview): Review response metrics and KPI summary widgets.",
             "Dashboard renders average response time (minutes), incidents by category (pie chart), and incidents by barangay (bar chart)."),
            ("Date Range Analytics Aggregation: Switch analytics filter between Today, This Week, This Month, and Custom Range.",
             "Chart datasets recalculate dynamically matching filtered historical incidents.")
        ]
    )

    # Phase 11
    add_phase(
        "Phase 11 — Reporting, Printables, and Data Export (PDF & Excel)",
        "Validate official PDF Incident Reports, Shift Summary exports, formatted Excel spreadsheets (.xlsx), and print stylesheets.",
        [
            ("Official Incident PDF Export: On Incident Detail page, click 'Export PDF'.",
             "PDF downloads formatted with official Tuguegarao City / RESQ-Link letterhead, incident summary, timeline, and signature block."),
            ("Shift Summary PDF Report: From History / Overview, generate daily Shift Summary PDF.",
             "Multi-page PDF compiles all incidents handled during the operational shift with category breakdowns and dispatcher sign-offs."),
            ("Export Incident History to Excel (.xlsx): In Incident History table, click 'Export to Excel'.",
             "Clean .xlsx file downloads with formatted headers, auto-fitted columns, dates, priorities, agencies, and resolution outcomes."),
            ("Export Civilian Registry to Excel: In Super Admin Civilians table, export civilian records to Excel.",
             "Spreadsheet exports full registered user dataset with verification statuses, registration dates, and contact numbers."),
            ("Export Equipment & Resource Status to Excel: In Resources table, export equipment inventory to Excel.",
             "Spreadsheet lists all vehicles, equipment readiness, assigned agencies, and maintenance statuses."),
            ("Browser Print View Optimization: Trigger browser print (Ctrl+P) on Incident Detail and Overview pages.",
             "Print stylesheet removes navigation sidebar, action buttons, and dark backgrounds, outputting clean, readable white documents.")
        ]
    )

    # Phase 12
    add_phase(
        "Phase 12 — Security, Role-Based Access Control, and Edge Cases",
        "Validate RBAC boundaries, unauthenticated route guards, offline-online sync, location permission denials, and multi-device concurrency.",
        [
            ("Role Route Protection (Dispatcher -> Admin): Attempt to access /admin/* while logged in as a Dispatcher.",
             "Access Denied page (/access-denied) rendered; Dispatcher cannot view or modify Super Admin settings."),
            ("Role Route Protection (Civilian -> Web Portal): Attempt to log into Web Console (/login) with Civilian credentials.",
             "Authentication rejected with role mismatch error; Civilian directed to use mobile application."),
            ("Unauthenticated Route Redirection: Attempt to navigate directly to /command-center/overview or /admin/users in incognito window.",
             "User redirected immediately to /login with original destination preserved in query redirect param."),
            ("Concurrent Incident Editing: Open the same active incident on two separate dispatcher browser windows simultaneously.",
             "Firestore real-time listeners synchronize status changes across both sessions without data collision or race conditions."),
            ("Mobile Network Disconnect & Reconnect: Put Responder mobile app in Airplane Mode while En Route, tap On Scene, then restore network.",
             "Status transition buffers locally and commits to Firestore immediately upon reconnect; zero application crashes."),
            ("GPS Permission Denied Handling: Deny location permission on Civilian mobile app during emergency reporting.",
             "App displays informative permission guidance banner and enables manual pinpointing on interactive map."),
            ("Mobile Background / Resume Lifecycle: Background Civilian and Responder apps for 5 minutes during active incident, then resume.",
             "Apps seamlessly re-establish Firestore snapshot listeners and presence connection without UI freezing or stale data."),
            ("Responsive Web Layout: Test Command Center and Admin web portals on Laptop (1366x768), Desktop (1920x1080), and Tablet viewports.",
             "Sidebar collapses into responsive hamburger menu; data tables scroll smoothly; map controls remain fully accessible."),
            ("Session Revocation / Deactivated User Token: Deactivate an active dispatcher user in Super Admin while that user is logged in.",
             "Deactivated user is logged out on next token refresh or operational action with 'Session Terminated' notice."),
            ("Emergency Panic Button Debounce: Rapidly tap the Civilian Panic SOS button 5+ times in 2 seconds.",
             "Debounce logic prevents multiple duplicate emergency records from being created; exactly one incident queued.")
        ]
    )

    # Sign-off Section
    p_sign_heading = doc.add_heading("Sign-off", level=1)
    p_sign_heading.paragraph_format.space_before = Pt(16)
    p_sign_heading.paragraph_format.space_after = Pt(4)
    for run in p_sign_heading.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x0F, 0x2D, 0x59)

    p_sign_desc = doc.add_paragraph()
    p_sign_desc.paragraph_format.space_before = Pt(0)
    p_sign_desc.paragraph_format.space_after = Pt(8)
    run_sd = p_sign_desc.add_run("All testing phases completed and verified. Every Fail item has been resolved, re-tested, or documented with an approved operational workaround.")
    run_sd.font.name = "Calibri"
    run_sd.font.size = Pt(9.5)
    run_sd.font.italic = True
    run_sd.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    sign_headers = ["Tester / Developer Name", "Phases Tested & Scope", "Signature", "Date"]
    sign_data = [
        ["", "", "", ""],
        ["", "", "", ""],
        ["", "", "", ""],
        ["", "", "", ""]
    ]
    sign_widths = [Inches(2.2), Inches(2.3), Inches(1.3), Inches(1.2)]
    create_styled_table(doc, sign_headers, sign_data, sign_widths)

    # Save
    doc.save(output_path)
    print(f"Successfully generated {output_path}")

if __name__ == "__main__":
    build_resq_checklist()
