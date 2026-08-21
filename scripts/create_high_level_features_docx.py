from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "RESQ-Link_High-Level_Features.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def add_feature_section(document, number, title, subtitle, features):
    heading = document.add_heading(f"{number}. {title}", level=1)
    heading.paragraph_format.space_before = Pt(12)
    intro = document.add_paragraph(subtitle)
    intro.style = document.styles["Subtitle"]
    for feature in features:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(feature)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.size = Pt(28)
styles["Title"].font.color.rgb = RGBColor(20, 83, 45)
styles["Title"].font.bold = True
styles["Subtitle"].font.name = "Aptos"
styles["Subtitle"].font.size = Pt(11)
styles["Subtitle"].font.color.rgb = RGBColor(71, 85, 105)
for style_name in ("Heading 1", "Heading 2"):
    styles[style_name].font.name = "Aptos Display"
    styles[style_name].font.color.rgb = RGBColor(20, 83, 45)

title = doc.add_paragraph()
title.style = doc.styles["Title"]
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("RESQ-Link")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("High-Level System Features")
run.bold = True
run.font.size = Pt(15)
run.font.color.rgb = RGBColor(51, 65, 85)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.add_run(
    "An integrated platform connecting civilians, command centers, and field responders throughout the emergency-response lifecycle."
).italic = True

doc.add_paragraph()

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
table.style = "Table Grid"
labels = ["Civilian App", "Dispatcher App", "Responder App", "Super Admin"]
for cell, label in zip(table.rows[0].cells, labels):
    set_cell_shading(cell, "DCFCE7")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor(20, 83, 45)

doc.add_paragraph()

add_feature_section(doc, 1, "Civilian Mobile App", "Emergency reporting and support for members of the public.", [
    "Secure civilian registration and login.",
    "Submit emergency reports with location and incident details.",
    "Support for fire, medical, vehicular, police, electrical, and other emergencies.",
    "Review and confirm reports before submission.",
    "Track emergency-report history and current status.",
    "View dispatched responders on a map.",
    "Place emergency voice calls and receive operational notifications.",
    "Request incident or surveillance footage and track request status.",
    "Manage profile, appearance, privacy, help, and support settings.",
    "Report application issues and access frequently asked questions.",
])

add_feature_section(doc, 2, "Dispatcher Command-Center App", "Central coordination, resource dispatch, monitoring, and reporting.", [
    "Monitor active and pending incidents through a real-time dashboard.",
    "Accept incident intake from the civilian app, calls, SMS, radio, walk-ins, or manual entry.",
    "Validate, classify, prioritize, assign, and dispatch incidents.",
    "Use an interactive map to monitor incidents and responder locations.",
    "Track incidents through assigned, en route, on-scene, and resolved stages.",
    "Manage emergency resources, vehicles, teams, and crew assignments.",
    "Match available resources with active responders or agencies.",
    "Recover, dispatch, or remove incidents awaiting resources.",
    "Use direct and group operational messaging with responders.",
    "Search and analyze incident history, response times, priorities, and outcomes.",
    "Manage incident categories, dispatch rules, footage requests, and dispatcher teams.",
])

add_feature_section(doc, 3, "Responder Mobile App", "Real-time field response and communication for emergency personnel.", [
    "Secure responder login with role and account-status verification.",
    "Receive assigned emergencies and updates in real time.",
    "View case priority, status, location, incident details, and reporter information.",
    "Use map-based incident and responder navigation.",
    "Update operational status to en route, on scene, or resolved.",
    "Communicate directly with dispatchers through operational messaging.",
    "Receive assignment and incident-status notifications.",
    "Access responder settings, location controls, and support pages.",
])

add_feature_section(doc, 4, "Super-Admin App", "Central administration of organizations, roles, and user accounts.", [
    "Manage dispatcher and field-responder accounts separately.",
    "Create and manage civilian and command-center accounts.",
    "Assign agency roles including BFP, PNP, MDRRMO, Ambulance, and PCG.",
    "Provision accounts securely through Firebase server-side services.",
    "Control role-based access and active-account status.",
])

doc.add_heading("Shared Platform Capabilities", level=1)
for feature in [
    "Firebase authentication and role-based authorization.",
    "Real-time Firestore synchronization across applications.",
    "Secure database and file-storage access rules.",
    "Geographic incident mapping and responder tracking.",
    "Coordinated incident lifecycle across civilians, dispatchers, resources, and responders.",
    "Media attachments, footage requests, operational messaging, and centralized reporting.",
]:
    doc.add_paragraph(feature, style="List Bullet")

doc.add_heading("End-to-End Workflow", level=1)
workflow = doc.add_paragraph()
workflow.alignment = WD_ALIGN_PARAGRAPH.CENTER
workflow_run = workflow.add_run(
    "Civilian Reporting  →  Command-Center Intake  →  Resource Dispatch  →  Responder Tracking and Communication  →  Incident Resolution and Reporting"
)
workflow_run.bold = True
workflow_run.font.size = Pt(11)
workflow_run.font.color.rgb = RGBColor(20, 83, 45)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("RESQ-Link — High-Level System Features").font.size = Pt(8)

doc.save(OUTPUT)
print(OUTPUT)
