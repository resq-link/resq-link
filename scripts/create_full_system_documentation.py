from __future__ import annotations

from datetime import date
from pathlib import Path
from tempfile import TemporaryDirectory

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "RESQ-Link_System_and_Technical_Documentation.docx"
LOGO = ROOT / "apps" / "dispatcher-web-app" / "public" / "branding" / "resq-link-logo.png"
ICON = ROOT / "apps" / "dispatcher-web-app" / "public" / "branding" / "resq-link-icon.png"
SCREENSHOTS = ROOT / "docs" / "screenshots"
DISPATCHER_SCREEN = SCREENSHOTS / "dispatcher-web-login-live.png"
SUPER_ADMIN_SCREEN = SCREENSHOTS / "super-admin-live.png"
CIVILIAN_SCREEN = SCREENSHOTS / "civilian-mobile-live.png"
RESPONDER_SCREEN = SCREENSHOTS / "responder-mobile-live.png"

PRIMARY = "20A384"
DARK = "103D3A"
ACCENT = "9AFF55"
SLATE = "344B57"
MUTED = "667985"
MINT = "EAF7F3"
PALE = "F4FAF8"
LINE = "CFE5DE"
WHITE = "FFFFFF"
AMBER = "F4B740"
RED = "D95555"
BLUE = "3A7CA5"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    if placeholder:
        value_run = paragraph.add_run(placeholder)
        value_run.font.color.rgb = rgb(MUTED)
    paragraph.add_run()._r.append(end)


def add_bottom_border(paragraph, color=LINE, size="10") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_table(doc: Document, headers, rows, widths=None, font_size=8.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, heading in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, DARK)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(heading))
        run.bold = True
        run.font.color.rgb = rgb(WHITE)
        run.font.size = Pt(font_size)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        set_row_cant_split(row)
        for i, value in enumerate(values):
            cell = row.cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_idx % 2:
                set_cell_shading(cell, PALE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(value))
            run.font.size = Pt(font_size)
            run.font.color.rgb = rgb(SLATE)
        if widths:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullets(doc: Document, items, level=0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        if isinstance(item, tuple):
            label, body = item
            r = p.add_run(label)
            r.bold = True
            p.add_run(body)
        else:
            p.add_run(item)


def add_numbered(doc: Document, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, kind="info") -> None:
    colors = {
        "info": (MINT, PRIMARY),
        "warning": ("FFF6DF", AMBER),
        "risk": ("FDECEC", RED),
        "dark": (DARK, ACCENT),
    }
    fill, accent = colors[kind]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.12)
    table.columns[1].width = Inches(6.9)
    set_cell_shading(table.cell(0, 0), accent)
    set_cell_shading(table.cell(0, 1), fill)
    set_cell_margins(table.cell(0, 0), 20, 20, 20, 20)
    set_cell_margins(table.cell(0, 1), 120, 180, 120, 180)
    p = table.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = rgb(WHITE if kind == "dark" else DARK)
    p2 = table.cell(0, 1).add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.05
    p2.runs[0].font.color.rgb = rgb(WHITE if kind == "dark" else SLATE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_cell_shading(table.cell(0, 0), "152D2B")
    set_cell_margins(table.cell(0, 0), 130, 170, 130, 170)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Cascadia Mono"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb("D8F3EA")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_section_title(doc: Document, part: str, title: str, subtitle: str) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(part.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(PRIMARY)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(14)
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = rgb(DARK)
    add_bottom_border(title_p, PRIMARY, "20")
    sub = doc.add_paragraph(subtitle)
    sub.paragraph_format.space_after = Pt(24)
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = rgb(MUTED)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(MUTED)


def font(size, bold=False):
    for name in ("C:/Windows/Fonts/aptos.ttf", "C:/Windows/Fonts/arial.ttf"):
        if Path(name).exists():
            return ImageFont.truetype(name, size=size)
    return ImageFont.load_default()


def bold_font(size):
    for name in ("C:/Windows/Fonts/aptos-bold.ttf", "C:/Windows/Fonts/arialbd.ttf"):
        if Path(name).exists():
            return ImageFont.truetype(name, size=size)
    return font(size, True)


def rounded_box(draw, xy, fill, outline, title, lines, title_color=DARK):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=24, fill=fill, outline=outline, width=3)
    resolved_title_color = f"#{title_color}" if len(title_color) == 6 and not title_color.startswith("#") else title_color
    draw.text((x1 + 24, y1 + 20), title, font=bold_font(27), fill=resolved_title_color)
    y = y1 + 65
    body_color = "#D8F3EA" if fill == "#103D3A" else "#344B57"
    for line in lines:
        draw.text((x1 + 25, y), line, font=font(19), fill=body_color)
        y += 28


def arrow(draw, start, end, color="#20A384", width=7):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    dx, dy = ex - sx, ey - sy
    if abs(dx) >= abs(dy):
        sign = 1 if dx > 0 else -1
        pts = [(ex, ey), (ex - 18 * sign, ey - 12), (ex - 18 * sign, ey + 12)]
    else:
        sign = 1 if dy > 0 else -1
        pts = [(ex, ey), (ex - 12, ey - 18 * sign), (ex + 12, ey - 18 * sign)]
    draw.polygon(pts, fill=color)


def make_architecture_diagram(path: Path) -> None:
    canvas = Image.new("RGB", (1600, 900), "#F4FAF8")
    draw = ImageDraw.Draw(canvas)
    draw.text((70, 35), "RESQ-Link Logical Architecture", font=bold_font(40), fill="#103D3A")
    draw.text((70, 88), "Four role-specific experiences share one operational data and service layer.", font=font(23), fill="#667985")

    rounded_box(draw, (70, 170, 450, 330), "#FFFFFF", "#20A384", "Civilian Mobile", ["Expo / React Native", "Reports • Tracking • Calls"])
    rounded_box(draw, (70, 390, 450, 550), "#FFFFFF", "#20A384", "Responder Mobile", ["Expo / React Native", "Assignments • GPS • Reports"])
    rounded_box(draw, (70, 610, 450, 770), "#FFFFFF", "#20A384", "Dispatcher Web", ["Next.js / React", "Intake • Dispatch • Analytics"])
    rounded_box(draw, (1150, 280, 1530, 440), "#FFFFFF", "#20A384", "Super Admin", ["Next.js + Admin SDK", "Accounts • KYC • Recovery"])

    rounded_box(draw, (600, 255, 1030, 500), "#EAF7F3", "#103D3A", "@packages/firebase", ["Typed domain services", "Auth • incidents • resources", "teams • messages • calls", "presence • storage • rules"])
    rounded_box(draw, (600, 610, 1030, 800), "#103D3A", "#103D3A", "Firebase Platform", ["Auth • Firestore • Storage", "Realtime DB • Admin SDK"], "#9AFF55")

    arrow(draw, (450, 250), (600, 330))
    arrow(draw, (450, 470), (600, 405))
    arrow(draw, (450, 690), (600, 475))
    arrow(draw, (1150, 360), (1030, 390))
    arrow(draw, (815, 500), (815, 610))
    draw.text((1085, 595), "External services", font=bold_font(22), fill="#103D3A")
    draw.text((1085, 635), "Google Maps / Mapbox", font=font(20), fill="#344B57")
    draw.text((1085, 670), "Agora voice", font=font(20), fill="#344B57")
    draw.text((1085, 705), "Resend email", font=font(20), fill="#344B57")
    draw.text((1085, 740), "Vercel / EAS", font=font(20), fill="#344B57")
    canvas.save(path, quality=95)


def make_workflow_diagram(path: Path) -> None:
    canvas = Image.new("RGB", (1600, 620), "#FFFFFF")
    draw = ImageDraw.Draw(canvas)
    draw.text((70, 35), "Emergency Response Operating Flow", font=bold_font(40), fill="#103D3A")
    steps = [
        ("1", "REPORT", "Civilian app, call, SMS,\nradio, walk-in, manual"),
        ("2", "TRIAGE", "Validate, classify,\nprioritize, enrich"),
        ("3", "ASSIGN", "Team, agency, resource,\nand primary responder"),
        ("4", "RESPOND", "Accept, en route,\non scene, communicate"),
        ("5", "CLOSE", "Assessment, post-report,\nresolution, analytics"),
    ]
    colors = ["#EAF7F3", "#DFF3ED", "#D4EFE6", "#C9EADF", "#BEE6D8"]
    x = 60
    for idx, (num, title, body) in enumerate(steps):
        y1, y2 = 170, 480
        draw.rounded_rectangle((x, y1, x + 260, y2), radius=26, fill=colors[idx], outline="#20A384", width=3)
        draw.ellipse((x + 94, y1 + 24, x + 166, y1 + 96), fill="#103D3A")
        tw = draw.textbbox((0, 0), num, font=bold_font(32))[2]
        draw.text((x + 130 - tw / 2, y1 + 39), num, font=bold_font(32), fill="#9AFF55")
        title_w = draw.textbbox((0, 0), title, font=bold_font(25))[2]
        draw.text((x + 130 - title_w / 2, y1 + 125), title, font=bold_font(25), fill="#103D3A")
        for line_idx, line in enumerate(body.split("\n")):
            line_w = draw.textbbox((0, 0), line, font=font(18))[2]
            draw.text((x + 130 - line_w / 2, y1 + 184 + line_idx * 30), line, font=font(18), fill="#344B57")
        if idx < len(steps) - 1:
            arrow(draw, (x + 260, 325), (x + 315, 325), width=6)
        x += 315
    canvas.save(path, quality=95)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(17)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = rgb(SLATE)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in (
        ("Title", 32, DARK, 0, 8),
        ("Subtitle", 13, MUTED, 0, 8),
        ("Heading 1", 21, DARK, 16, 7),
        ("Heading 2", 14.5, PRIMARY, 12, 5),
        ("Heading 3", 11.5, DARK, 9, 3),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = style_name != "Subtitle"
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        styles[style_name].font.name = "Aptos"
        styles[style_name].font.size = Pt(9.3)
        styles[style_name].font.color.rgb = rgb(SLATE)

    for sec in doc.sections:
        header = sec.header
        table = header.add_table(rows=1, cols=2, width=Inches(6.7))
        table.autofit = False
        table.columns[0].width = Inches(4.9)
        table.columns[1].width = Inches(1.8)
        left = table.cell(0, 0).paragraphs[0]
        left.paragraph_format.space_after = Pt(0)
        r = left.add_run("RESQ-LINK  /  SYSTEM DOCUMENTATION")
        r.bold = True
        r.font.size = Pt(7.5)
        r.font.color.rgb = rgb(DARK)
        right = table.cell(0, 1).paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right.paragraph_format.space_after = Pt(0)
        if ICON.exists():
            right.add_run().add_picture(str(ICON), width=Inches(0.22))
        add_bottom_border(header.paragraphs[0], PRIMARY, "8")

        footer = sec.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.add_run("Internal Technical Reference   •   Version 1.1   •   ").font.size = Pt(7.5)
        add_field(p, "PAGE", "1")
        p.add_run(" of ").font.size = Pt(7.5)
        add_field(p, "NUMPAGES", "1")
        for run in p.runs:
            run.font.color.rgb = rgb(MUTED)
            run.font.size = Pt(7.5)

    settings = doc.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "RESQ-Link System and Technical Documentation"
    props.subject = "High-level product overview and full technical reference"
    props.author = "RESQ-Link Engineering Team"
    props.keywords = "RESQ-Link, emergency response, architecture, Firebase, Expo, Next.js"
    props.comments = "Generated from repository evidence at commit eef022d."

    with TemporaryDirectory(prefix="resq_doc_") as temp_dir:
        temp = Path(temp_dir)
        architecture = temp / "architecture.png"
        workflow = temp / "workflow.png"
        make_architecture_diagram(architecture)
        make_workflow_diagram(workflow)

        # Cover
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(55)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(4.65))
        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.paragraph_format.space_before = Pt(20)
        line.paragraph_format.space_after = Pt(14)
        line.add_run("SYSTEM & TECHNICAL DOCUMENTATION").bold = True
        line.runs[0].font.size = Pt(11)
        line.runs[0].font.color.rgb = rgb(PRIMARY)
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style = "Title"
        title.add_run("Integrated Emergency Response Platform")
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.style = "Subtitle"
        subtitle.add_run("Executive overview, operating model, solution architecture, data design, security, deployment, and engineering reference")
        doc.add_paragraph()
        cover_table = doc.add_table(rows=2, cols=3)
        cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [("VERSION", "1.1"), ("BASELINE", "eef022d"), ("DATE", "20 August 2026")]
        for col, (label, value) in enumerate(labels):
            c1 = cover_table.cell(0, col)
            c2 = cover_table.cell(1, col)
            set_cell_shading(c1, DARK)
            set_cell_shading(c2, MINT)
            for c in (c1, c2):
                set_cell_margins(c, 100, 150, 100, 150)
                c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = c1.paragraphs[0].add_run(label)
            r1.bold = True
            r1.font.size = Pt(7.5)
            r1.font.color.rgb = rgb(ACCENT)
            r2 = c2.paragraphs[0].add_run(value)
            r2.bold = True
            r2.font.size = Pt(10)
            r2.font.color.rgb = rgb(DARK)
        statement = doc.add_paragraph()
        statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
        statement.paragraph_format.space_before = Pt(55)
        statement.add_run("Connecting civilians, command centers, and field responders across the complete emergency lifecycle.").italic = True
        statement.runs[0].font.color.rgb = rgb(MUTED)
        statement.runs[0].font.size = Pt(10)
        source = doc.add_paragraph()
        source.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source.paragraph_format.space_before = Pt(40)
        source.add_run("Prepared from the current RESQ-Link monorepo and public branding resources.").font.size = Pt(8)

        # Document control and TOC
        doc.add_page_break()
        doc.add_heading("Document Control", level=1)
        add_table(doc, ["Field", "Value"], [
            ("Document title", "RESQ-Link System and Technical Documentation"),
            ("Purpose", "Provide one authoritative high-level and technical reference for stakeholders, implementers, operators, and maintainers."),
            ("Scope", "Civilian mobile, responder mobile, dispatcher web, super-admin web, shared Firebase package, integrations, security rules, and deployment configuration."),
            ("Repository baseline", "Local main at commit eef022d, dated 20 August 2026"),
            ("Brand source", "apps/dispatcher-web-app/public/branding/resq-link-logo.png"),
            ("Document owner", "RESQ-Link Engineering Team"),
            ("Classification", "Internal technical reference; review before external distribution"),
        ], widths=[1.6, 5.3], font_size=8.7)
        doc.add_heading("Revision History", level=2)
        add_table(doc, ["Version", "Date", "Change", "Owner"], [
            ("1.1", "20 Aug 2026", "Added verified live-runtime application screenshots and capture provenance", "RESQ-Link Engineering Team"),
            ("1.0", "20 Aug 2026", "Initial consolidated high-level and technical documentation", "RESQ-Link Engineering Team"),
        ], widths=[0.8, 1.1, 4.0, 1.3])
        doc.add_heading("How to Use This Document", level=2)
        add_bullets(doc, [
            ("Executives and project stakeholders — ", "read Parts I and II for scope, value, users, capabilities, and operating flow."),
            ("Engineering and architecture teams — ", "use Parts III through VII for component, data, integration, security, and deployment detail."),
            ("Operations and support teams — ", "use Parts VIII and IX for setup, verification, troubleshooting, and continuity guidance."),
            ("Auditors and reviewers — ", "use the appendices for routes, environment variables, terminology, and source paths."),
        ])
        doc.add_page_break()
        doc.add_heading("Table of Contents", level=1)
        toc = doc.add_paragraph()
        add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "Right-click and select Update Field in Microsoft Word")
        add_callout(doc, "Navigation note", "The table of contents is configured to refresh automatically when the document opens in Microsoft Word. If page numbers are blank, select the table and update the field.")

        # PART I
        add_section_title(doc, "Part I", "Executive and High-Level Overview", "What RESQ-Link is, who it serves, and how the platform improves coordinated emergency response.")
        doc.add_heading("1. Executive Summary", level=1)
        doc.add_paragraph(
            "RESQ-Link is an integrated emergency-response platform designed to connect the public, command-center personnel, field responders, and system administrators through one coordinated digital workflow. It combines two Expo mobile applications, two Next.js web applications, and a shared Firebase service package in a single npm-workspace monorepo."
        )
        doc.add_paragraph(
            "The platform supports the full operational chain: authenticated civilian reporting, command-center intake and triage, incident creation and prioritization, team and resource dispatch, responder acceptance and status updates, live communication, post-incident reporting, historical analysis, and administrative governance. Real-time Firebase subscriptions keep role-specific interfaces synchronized while shared domain modules reduce duplication and enforce common lifecycle rules."
        )
        add_callout(doc, "Core value proposition", "A report entered once becomes a shared operational record that can be validated, elevated into an incident, assigned to a team and resources, acted on by responders, and closed with traceable outcomes.", "dark")

        doc.add_heading("1.1 Strategic Outcomes", level=2)
        add_bullets(doc, [
            "Shorter information handoffs between civilians, dispatchers, and responders.",
            "A shared operational picture for incidents, resources, teams, priorities, and locations.",
            "Consistent incident and emergency-report lifecycle handling across applications.",
            "Role-based accountability through timestamps, assignments, status transitions, acknowledgments, and post-incident records.",
            "Centralized administration of civilian, dispatcher, command-center, and responder accounts.",
            "Reusable platform services for authentication, storage, messaging, calls, presence, reporting, and real-time synchronization.",
        ])
        doc.add_heading("1.2 Primary Stakeholders", level=2)
        add_table(doc, ["Stakeholder", "Primary need", "RESQ-Link capability"], [
            ("Civilian / reporter", "Fast, trustworthy emergency reporting and visibility", "Guided reports, location, attachments, status history, map tracking, voice calls, and account support"),
            ("Command-center dispatcher", "Accurate triage and coordinated deployment", "Real-time intake, incident management, team/resource assignment, map, messaging, alerting, and reporting"),
            ("Field responder", "Clear assignments and safe operational coordination", "Incident details, navigation, accept/decline, en-route/on-scene updates, communication, and post-reporting"),
            ("Super administrator", "Governance and secure account lifecycle", "Account provisioning, KYC review, email verification support, recovery APIs, and role control"),
            ("Engineering / operations", "Maintainable and deployable platform", "Monorepo, shared typed services, security rules, environment conventions, Vercel/EAS configuration, and scripts"),
        ], widths=[1.45, 2.2, 3.3])

        doc.add_heading("2. Platform at a Glance", level=1)
        add_table(doc, ["Product", "Audience", "Technology", "Primary responsibilities"], [
            ("Civilian Mobile App", "Residents and reporters", "Expo 54, React Native 0.81, Expo Router", "Registration, email verification/KYC states, reporting, report history, emergency tracking map, calls, profile and settings"),
            ("Dispatcher Web App", "Command centers and dispatchers", "Next.js 15, React 19, TypeScript, Leaflet", "Intake, triage, incidents, teams, resources, operational messaging, voice coordination, maps, history, analytics, exports"),
            ("Responder Mobile App", "BFP, PNP, MDRRMO, ambulance, PCG and field teams", "Expo 54, React Native 0.81, Expo Router", "Assignments, case detail, map/navigation, availability/presence, status transitions, alerts, communication, scene/post reports"),
            ("Super-Admin Web App", "Platform administrators", "Next.js 15, React 19, Firebase Admin SDK", "User provisioning, command-center management, KYC approval/rejection, OTP support, password recovery"),
            ("Shared Firebase Package", "All applications", "TypeScript, Firebase Web/Admin SDKs", "Typed domain operations, synchronization, authentication, storage, status logic, presence, messaging, calls and security assets"),
        ], widths=[1.25, 1.45, 1.7, 2.75], font_size=8)
        doc.add_picture(str(architecture), width=Inches(6.95))
        add_caption(doc, "Figure 1. Logical architecture and shared service boundary.")

        doc.add_heading("3. End-to-End Operating Model", level=1)
        doc.add_picture(str(workflow), width=Inches(6.95))
        add_caption(doc, "Figure 2. Emergency response operating flow.")
        add_numbered(doc, [
            "Report or intake — a civilian submits through the mobile app, or the command center records a call, SMS, walk-in, radio, or manual source.",
            "Triage — dispatch personnel validate the information, determine category/subtype, priority, required agencies, location, and follow-up details.",
            "Incident formation — eligible emergency reports may be linked or elevated into a canonical incident, preserving report references and field evidence.",
            "Assignment and dispatch — the command center assigns the active operational team, agencies, resources, and responders; dispatch records preserve the relationship.",
            "Field response — responders accept or decline, navigate to the scene, update en-route and on-scene states, communicate, and submit assessments.",
            "Resolution — post-incident information, response timing, outcome, resource release, history movement, and analytics complete the lifecycle.",
        ])
        doc.add_heading("3.1 Operational Principles", level=2)
        add_bullets(doc, [
            ("Single operational thread — ", "reports, incidents, resources, dispatch records, teams, and call/chat sessions retain cross-references."),
            ("Permanent incident ownership — ", "assigned team identifiers and snapshots remain on the incident for filtering, history, and accountability."),
            ("Real-time by default — ", "Firestore listeners update active lists and detail views; Realtime Database improves responder presence accuracy."),
            ("Role-aware actions — ", "security rules and UI gates constrain account, incident, messaging, storage, and administrative operations."),
            ("Graceful degradation — ", "query fallbacks, map-key checks, in-memory sorting, and Firestore presence fallback reduce hard failures."),
        ])

        doc.add_heading("4. High-Level Capability Catalog", level=1)
        capability_rows = [
            ("Identity and onboarding", "Email/password authentication, civilian registration, email OTP, KYC workflow, role/designation checks, password recovery"),
            ("Emergency reporting", "Structured incident types, GPS/manual location, field assessment, people involved, photos, review and submission"),
            ("Command-center intake", "Multiple source channels, categorization, priority, additional details, report linking, incident elevation"),
            ("Dispatch operations", "Operational team on duty, agency routing, resource matching, assignments, reassignment, and dispatch records"),
            ("Field response", "Assignment subscription, accept/decline, en route, touchdown/on scene, responder assessment, post-incident report"),
            ("Situational awareness", "Incident maps, reporter/responder coordinates, quadrant data, map filters, markers, banners, and status visualization"),
            ("Communication", "Direct/group chat, incident voice-call sessions, Agora token issuance, alerts, acknowledgments, and notifications"),
            ("Resource management", "Vehicles/equipment, availability, station/current coordinates, responder binding, maintenance/offline states"),
            ("Evidence and records", "Emergency photos, post-report photos, footage requests, history, report exports, PDF/Excel support"),
            ("Governance", "Account management, KYC decisioning, command-center and responder provisioning, server-only OTP records"),
            ("Operational analytics", "Priority/status summaries, response time, team reporting, date filters, print/export workflows"),
            ("Platform resilience", "Lazy Firebase initialization, auth readiness, missing-index fallback, presence fallback, map configuration fallback"),
        ]
        add_table(doc, ["Capability domain", "Implemented scope"], capability_rows, widths=[1.8, 5.2], font_size=8.4)

        # PART II
        add_section_title(doc, "Part II", "Application Portfolio", "Role-specific experiences, major routes, responsibilities, and boundaries.")
        doc.add_heading("5. Civilian Mobile Application", level=1)
        doc.add_paragraph("The civilian application provides the public-facing entry point into RESQ-Link. Its feature-oriented architecture separates authentication, dashboard, emergency reporting, incident tracking, history, profile, settings, and voice-call concerns behind thin Expo Router route files.")
        doc.add_heading("5.1 Functional Scope", level=2)
        add_bullets(doc, [
            "Registration with personal details, government-ID type and KYC evidence; email verification and pending/rejected/active account states.",
            "Forgot-password request, OTP validation and password reset flows.",
            "Guided emergency report creation with incident type/profile, narrative fields, people involved, location, photos, review, and submission overlay.",
            "GPS acquisition with manual address/pin support and a safe fallback when a Google Maps key is unavailable in a native Android build.",
            "Dashboard, report history, live incident tracking, status timeline, responder/map context, emergency confirmation and voice-call entry points.",
            "Profile, appearance, notification preferences, privacy/security, FAQs, help/support, and issue reporting.",
        ])
        doc.add_heading("5.2 Architecture", level=2)
        add_table(doc, ["Layer", "Key locations", "Responsibility"], [
            ("Routes", "src/app/(auth), (main), (settings)", "Navigation groups and route-to-screen adapters"),
            ("Features", "src/features/*", "Feature screens, components, hooks, constants and utilities"),
            ("Shared UI", "src/components/*", "Reusable controls, navigation, feedback and badges"),
            ("State", "src/stores/*", "Zustand authentication and user state"),
            ("Hooks/theme", "src/hooks/*; src/theme/*", "Cross-feature behavior, theme provider, palettes and factories"),
            ("Services", "src/services/*; @packages/firebase", "API, Agora voice and Firebase domain access"),
            ("Configuration", "app.json; app.config.js; metro.config.js", "Expo identity, plugins, native permissions, environment injection and monorepo resolution"),
        ], widths=[1.2, 2.35, 3.4])

        doc.add_heading("6. Dispatcher Command-Center Web Application", level=1)
        doc.add_paragraph("The dispatcher application is the operational coordination center. It uses the Next.js App Router and a component-rich dashboard to turn incoming reports into managed incidents, assign operational ownership, monitor live status, and produce auditable reports.")
        doc.add_heading("6.1 Functional Scope", level=2)
        add_bullets(doc, [
            "Authenticated command-center login and protected navigation.",
            "Overview and active-incident monitoring with priority alerts and operational widgets.",
            "Multi-source intake, report detail, additional-information handling, triage, incident creation, association, and deduplication support.",
            "Incident management with priority, status, agency routing, assignment, team-on-duty context, touch-down timing, history and outcome records.",
            "Resource and team administration, including assigned responders, current location, availability, and operational team filtering.",
            "Interactive incident/resource maps using Leaflet and configurable Mapbox tiles/styles.",
            "Direct/group operational messaging, incident voice calls through Agora, notification/alarm behavior, and assistant endpoint integration.",
            "Reports, historical filtering, response analytics, print views, PDF export and Excel export.",
        ])
        doc.add_heading("6.2 Major Pages", level=2)
        add_table(doc, ["Route", "Purpose"], [
            ("/overview", "Command-center summary and live operational widgets"),
            ("/intake", "Incoming civilian/manual intake and triage workspace"),
            ("/incidents", "Active incident list and management"),
            ("/incident-management", "Incident rules, categories and management workflows"),
            ("/map", "Geographic operational view"),
            ("/resources", "Vehicles, equipment and responder/resource bindings"),
            ("/teams", "Team records, membership and on-duty coordination"),
            ("/history", "Resolved/unresolved operational history"),
            ("/report and /report/incidents", "Analytics, filters, print and export outputs"),
            ("/footage-requests", "Evidence/CCTV request administration"),
        ], widths=[2.0, 5.0])

        doc.add_heading("7. Responder Mobile Application", level=1)
        doc.add_paragraph("The responder application is optimized for field use. It subscribes responders to assigned reports/incidents, exposes concise case details and map context, and writes operational status and assessment updates back to the shared data layer.")
        doc.add_heading("7.1 Functional Scope", level=2)
        add_bullets(doc, [
            "Email/password login with active account, role, and responder-designation validation.",
            "Dashboard cards for assignments, priority, status and immediate actions.",
            "Case detail view with reporter data, location, media, narrative/field assessment, timeline, team/resource context and communication controls.",
            "Accept/decline assignment, mark en route, record touchdown/on-scene state, complete response, and submit post-incident or scene-assessment information.",
            "Map explorer for case and responder context, with a configuration-aware native Google Maps fallback.",
            "Notifications, haptics, alarm behavior, operational chat/calls, settings, location controls, about and support screens.",
            "Realtime Database presence with Firestore online-status fallback when RTDB is not configured.",
        ])
        doc.add_heading("7.2 Route Model", level=2)
        add_table(doc, ["Route group", "Routes", "Responsibility"], [
            ("Authentication", "/(auth)/login; /index", "Session gate and role-aware login"),
            ("Primary tabs", "/(tabs)/dashboard, map, notifications, settings", "Daily responder workspace"),
            ("Incident detail", "/incident/[id]", "Assignment detail, status transitions and reporting"),
            ("Support", "/support/about, help-support, location", "Reference, support and device/location settings"),
        ], widths=[1.4, 2.7, 2.9])

        doc.add_heading("8. Super-Admin Web Application", level=1)
        doc.add_paragraph("The super-admin application governs identities and account status. Browser-side Firebase authentication is combined with server-side Firebase Admin SDK routes for privileged account and OTP operations.")
        add_bullets(doc, [
            "Manage responders/dispatchers by agency role and designation.",
            "Manage civilian and command-center accounts.",
            "Review civilian KYC records and approve or reject with a recorded reason.",
            "Send and verify email OTP records using server-only Firestore collections.",
            "Issue and validate password-reset OTPs and update credentials through privileged APIs.",
            "Use Firebase service-account credentials only on the server side; use Resend for transactional email delivery.",
        ])
        add_table(doc, ["Administrative API", "Responsibility"], [
            ("POST /api/create-dispatcher", "Create dispatcher/agency account and profile"),
            ("POST /api/create-responder", "Create field-responder account"),
            ("POST /api/create-civilian", "Create civilian identity/profile"),
            ("POST /api/create-command-center", "Create command-center identity/profile"),
            ("POST /api/email-otp/send and /verify", "Email verification workflow"),
            ("POST /api/auth/forgot-password/send and /reset", "Password-recovery OTP workflow"),
            ("POST /api/kyc/approve and /reject", "Civilian KYC decision workflow"),
        ], widths=[3.0, 4.0])

        # LIVE APPLICATION GALLERY
        add_section_title(doc, "Visual Evidence", "Live Application Gallery", "Verified browser and Android-emulator captures from the local RESQ-Link runtime.")
        add_callout(
            doc,
            "Capture provenance",
            "Captured on 20 August 2026 from locally running application builds. Web views were rendered in Microsoft Edge headless mode; mobile views were captured from the Android emulator framebuffer. No synthetic UI, substituted screens, or mock credentials were used. Authentication-protected portals are shown at their genuine entry states.",
        )

        doc.add_heading("Dispatcher Command-Center Web", level=1)
        doc.add_paragraph("Live browser capture of the dispatcher portal's branded authentication entry point.")
        if DISPATCHER_SCREEN.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(DISPATCHER_SCREEN), width=Inches(6.9))
            caption = doc.add_paragraph("Figure 3. Dispatcher command-center web login — Microsoft Edge, local Next.js runtime.")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.style = "Caption"

        doc.add_page_break()
        doc.add_heading("Super-Admin Web", level=1)
        doc.add_paragraph("Live browser capture of the privileged administration portal's authentication entry point.")
        if SUPER_ADMIN_SCREEN.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(SUPER_ADMIN_SCREEN), width=Inches(6.9))
            caption = doc.add_paragraph("Figure 4. Super-admin web login — Microsoft Edge, local Next.js runtime.")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.style = "Caption"

        doc.add_page_break()
        doc.add_heading("Mobile Applications", level=1)
        doc.add_paragraph("Live Android captures showing the role-specific civilian and responder experiences rendered by their native application builds.")
        mobile_table = doc.add_table(rows=2, cols=2)
        mobile_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        mobile_table.autofit = False
        for column in mobile_table.columns:
            column.width = Inches(3.35)
        mobile_items = [
            (CIVILIAN_SCREEN, "Figure 5. Civilian mobile dashboard — live Android application state."),
            (RESPONDER_SCREEN, "Figure 6. Responder portal login — live Android development build."),
        ]
        for column_index, (image_path, caption_text) in enumerate(mobile_items):
            image_cell = mobile_table.cell(0, column_index)
            image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            image_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if image_path.exists():
                image_cell.paragraphs[0].add_run().add_picture(str(image_path), width=Inches(2.85))
            caption_cell = mobile_table.cell(1, column_index)
            caption_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            caption_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_cell.paragraphs[0].add_run(caption_text)
            caption_run.italic = True
            caption_run.font.size = Pt(8)
            caption_run.font.color.rgb = rgb(MUTED)

        # PART III
        add_section_title(doc, "Part III", "Solution Architecture", "Technology stack, monorepo structure, service boundaries, and runtime interaction patterns.")
        doc.add_heading("9. Technology Stack", level=1)
        add_table(doc, ["Layer", "Technology", "Use in RESQ-Link"], [
            ("Mobile", "Expo 54, React Native 0.81, React 19, Expo Router 6", "Civilian and responder applications; Android and iOS configuration"),
            ("Web", "Next.js 15, React 19, TypeScript", "Dispatcher and super-admin applications with App Router and API routes"),
            ("Domain/data", "Firebase 12 Web SDK; Firebase Admin SDK", "Authentication, Firestore, Storage, Realtime Database, privileged account operations"),
            ("State", "React hooks, contexts, Zustand, TanStack Query", "Local/global state, cached asynchronous state and screen-level orchestration"),
            ("Maps", "react-native-maps / Google Maps; Leaflet / Mapbox", "Native location and web operational mapping"),
            ("Communication", "Agora RTC; Firestore messaging; Expo notifications", "Voice calls, operational chat and device/user alerts"),
            ("UI", "Tailwind CSS, Lucide icons, bottom sheets, Reanimated/Moti", "Responsive web/mobile presentation and interaction"),
            ("Reporting", "jsPDF, jsPDF AutoTable, xlsx", "Command-center PDF, print and spreadsheet exports"),
            ("Deployment", "Vercel; Expo/EAS", "Web deployments and mobile build profiles"),
            ("Tooling", "npm workspaces, TypeScript, patch-package", "Monorepo dependency management, compilation and targeted patches"),
        ], widths=[1.2, 2.45, 3.35], font_size=8.3)

        doc.add_heading("10. Monorepo Structure", level=1)
        add_code(doc, """resq-link/
├── apps/
│   ├── civilian-mobile-app/      # Public Expo mobile experience
│   ├── dispatcher-web-app/       # Command-center Next.js application
│   ├── responder-mobile-app/     # Field responder Expo application
│   └── super-admin-web-app/      # Administrative Next.js application
├── packages/
│   └── firebase/                 # Shared typed Firebase/domain package
├── patches/                      # Root dependency patches
├── scripts/                      # Documentation and utility scripts
├── package.json                  # npm workspace and shared overrides
└── package-lock.json             # Workspace dependency lock""")
        doc.add_heading("10.1 Workspace Design", level=2)
        add_bullets(doc, [
            "The root package defines npm workspaces for apps/* and packages/* and pins cross-workspace React, Expo, React Native and tool versions through overrides.",
            "Each application consumes @packages/firebase through a local file workspace reference.",
            "The Firebase package compiles TypeScript from src/ to dist/ and exposes its main, admin and civilian-auth entry points.",
            "Vercel builds install the relevant web workspace plus @packages/firebase, build the shared package first, then build the selected web app.",
            "Metro configurations provide mobile resolution of the shared package in the monorepo.",
        ])

        doc.add_heading("11. Runtime Component Model", level=1)
        add_table(doc, ["Component", "Boundary", "Key collaborators"], [
            ("Civilian presentation/domain", "Feature screens, hooks and route adapters", "Shared Firebase functions, Expo device APIs, Maps, Agora"),
            ("Responder presentation/domain", "Modules for auth, dashboard, incidents, maps, notifications and settings", "Shared Firebase functions, RTDB presence, Maps, Agora"),
            ("Dispatcher presentation/domain", "App Router pages, contexts, operational/reporting components", "Firestore listeners, shared domain functions, Leaflet/Mapbox, Agora API"),
            ("Admin web/API", "Protected pages plus server-side API routes", "Firebase Admin, Firestore server access, Resend"),
            ("Shared Firebase domain", "Framework-neutral TypeScript modules", "Firebase Auth, Firestore, Storage, RTDB"),
            ("Security policy", "firestore.rules, storage.rules, database.rules.json", "Authenticated identities, role profiles, participant lists and ownership checks"),
        ], widths=[1.8, 2.45, 2.75])

        doc.add_heading("12. Shared Firebase Package", level=1)
        doc.add_paragraph("@packages/firebase acts as the system's primary domain and persistence boundary. Applications import typed operations instead of recreating Firestore queries and lifecycle rules independently.")
        module_rows = [
            ("config.ts", "Lazy app/auth/Firestore/Storage/RTDB initialization, Expo/Next environment resolution, auth readiness"),
            ("auth.ts / civilian-auth.ts", "Account creation, sign-in, role verification, registration and civilian account states"),
            ("emergencies.ts", "Civilian reports, subscriptions, assignment, linkage, follow-up details, status and post-report operations"),
            ("incidents.ts", "Canonical incidents, rule resolution, report elevation, dispatch, responder lifecycle and resource release"),
            ("priority.ts / alertAcknowledgment.ts", "Priority normalization, visual/audio behavior, escalation phases and acknowledgments"),
            ("incidentLifecycle.ts / incidentStatusVisual.ts", "Shared live/resolved predicates and normalized visual status tokens"),
            ("civilianFieldAssessment.ts", "Structured civilian narrative and follow-up field definitions/mapping"),
            ("responderAssessment.ts", "Scene-assessment definitions, parsing and submission"),
            ("resources.ts / teams.ts / operationalTeams.ts", "Resource CRUD, team records, assignment snapshots and matching"),
            ("commandCenterShift.ts", "Current team-on-duty persistence and subscription"),
            ("messaging.ts", "Direct/group threads, participants, messages and subscriptions"),
            ("callSessions.ts", "Incident voice-call session state machine and subscriptions"),
            ("dispatchers.ts / responderPresence.ts", "Location/online fields and RTDB presence lifecycle"),
            ("footageRequests.ts / storage.ts", "Evidence requests and image uploads"),
            ("quadrants.ts", "Operational quadrant definitions and barangay mapping"),
            ("admin.ts", "Server-only account provisioning and token verification"),
        ]
        add_table(doc, ["Module", "Primary responsibility"], module_rows, widths=[2.25, 4.75], font_size=8.1)

        # PART IV
        add_section_title(doc, "Part IV", "Data and Lifecycle Architecture", "Canonical records, relationships, statuses, synchronization, and operational ownership.")
        doc.add_heading("13. Data Stores", level=1)
        add_table(doc, ["Store", "Purpose", "Characteristics"], [
            ("Cloud Firestore", "Primary operational system of record", "Document collections, role-aware security rules, real-time subscriptions, indexed queries and history"),
            ("Firebase Authentication", "Identity and credential authority", "Email/password and phone-related SDK support; token verification for server APIs"),
            ("Firebase Storage", "KYC, emergency and post-report media", "Path-scoped rules for ownership, authenticated access and administrative review"),
            ("Firebase Realtime Database", "Responder online presence", "presence/responders/{uid}, .info/connected and onDisconnect cleanup; Firestore fallback"),
            ("Local device storage", "Session-supporting mobile state", "AsyncStorage/SecureStore used by mobile dependencies and app state"),
        ], widths=[1.7, 2.5, 2.8])

        doc.add_heading("14. Firestore Collection Catalog", level=1)
        collections = [
            ("admins", "Super-admin profiles", "Super-admin UI and server authorization"),
            ("users", "Civilian profiles, account status and KYC metadata", "Civilian app and super-admin KYC/account management"),
            ("dispatchers", "Dispatcher/responder profiles, roles, designation and location/online fields", "Dispatcher, responder and admin apps"),
            ("commandCenters", "Command-center profiles and current team-on-duty state", "Dispatcher authentication and operational context"),
            ("emergencies", "Original civilian/manual emergency reports and report-level lifecycle", "Civilian, dispatcher and responder experiences"),
            ("incidents", "Canonical command-center incident records", "Dispatch, response lifecycle, outcomes and analytics"),
            ("incidents/{id}/teamAssignmentHistory", "Team reassignment audit trail", "Operational history and reporting"),
            ("incidentTypeRules", "Subtype, priority, agency and resource-routing rules", "Incident creation and administration"),
            ("resources", "Vehicles/equipment, status, location, team and responder bindings", "Resource management and dispatch"),
            ("incidentDispatches", "Incident-to-agency/resource assignment records", "Traceable dispatch relationships"),
            ("teams", "Operational team catalog and membership context", "Team filtering, assignment and reporting"),
            ("chatThreads / messages", "Direct/group operational conversations", "Dispatcher-responder messaging"),
            ("callSessions", "Incident voice-call state and participants", "Civilian/responder calls with Agora channels"),
            ("footageRequests", "CCTV/evidence request lifecycle", "Civilian submission and command-center processing"),
            ("emailOtps", "Server-only email verification OTP records", "Super-admin API routes only"),
            ("passwordResetOtps", "Server-only password-reset OTP records", "Super-admin API routes only"),
        ]
        add_table(doc, ["Collection", "Data responsibility", "Primary consumers"], collections, widths=[1.75, 3.0, 2.25], font_size=7.9)

        doc.add_heading("15. Core Record Relationships", level=1)
        add_bullets(doc, [
            ("Emergency report → incident — ", "emergencies.incidentId links a reporter-originated record to a canonical incident; incidents.associatedReportIds supports the reverse relationship."),
            ("Report grouping — ", "primaryReportId enables related reports to be grouped before or alongside incident elevation."),
            ("Incident → dispatch — ", "incidentDispatches reference the incident, agency, resource and responder assignments."),
            ("Incident → resource — ", "incidents.assignedResourceIds and resources.assignedIncidentId represent current allocation; lifecycle functions synchronize release."),
            ("Incident → team — ", "assignedTeamId/name/code provide a permanent snapshot, while team history records capture reassignment."),
            ("Incident → call/chat — ", "call sessions carry incident identifiers; chat threads carry participant lists and operational conversation state."),
            ("User profiles → authentication — ", "Firebase Auth UID is the profile document key for users, dispatchers, command centers and admins."),
        ])

        doc.add_heading("16. Status and Lifecycle Models", level=1)
        add_table(doc, ["Domain", "States", "Notes"], [
            ("Civilian account", "pending_email_verification → pending_kyc_review → active; rejected", "Controls onboarding gate and administrative review"),
            ("Emergency report", "pending, linked, active, enroute, on_scene, done, resolved", "Supports legacy and current states while remaining synchronized with incident activity"),
            ("Incident", "new, awaiting_resources, liaison_pending, dispatched, enroute, on_scene, resolved, unresolved", "Canonical command-center lifecycle"),
            ("Resolution", "open, resolved, unresolved", "Separates outcome from operational status"),
            ("Resource", "available, assigned, en_route, on_scene, maintenance, offline", "Availability and deployment lifecycle"),
            ("Call session", "ringing, accepted, connected, ended, missed, failed", "Voice signaling state; media handled by Agora"),
            ("Footage request", "Defined in footageRequests domain", "Submitted and processed through civilian/command-center views"),
        ], widths=[1.45, 3.5, 2.05], font_size=8.1)
        add_callout(doc, "Lifecycle invariant", "Incident updates propagate to linked emergency reports where appropriate, while resource assignments and release operations maintain operational consistency. UI code should call shared lifecycle functions rather than writing status fields directly.")

        doc.add_heading("17. Priority, Escalation, and Alerting", level=1)
        add_bullets(doc, [
            "Priority values and rankings are centralized in priority.ts, including incident-type defaults and visual/map tokens.",
            "Acknowledgment fields include actor, timestamp, dispatcher identity, escalation level and alert timestamps.",
            "Shared helpers decide whether an alert requires forced visibility, acknowledgment UI, repeating audio, or escalation.",
            "Dispatcher contexts/components coordinate critical alerts and alarm unlocking; mobile clients use notification and haptic capabilities.",
            "Status-visual helpers normalize operational states to common labels, colors, Tailwind classes and pulse behavior.",
        ])

        doc.add_heading("18. Real-Time Synchronization", level=1)
        add_table(doc, ["Subscription", "Consumers", "Behavior"], [
            ("Emergency reports", "Civilian history/map; dispatcher intake; responder assignments", "Firestore onSnapshot, status filtering, sorted snapshots and metadata"),
            ("Incidents", "Dispatcher and responder", "Filtered real-time incident lists and detail updates"),
            ("Resources and teams", "Dispatcher", "Live resource/team availability and management"),
            ("Messages and threads", "Dispatcher and responder", "Participant-scoped thread/message listeners"),
            ("Call sessions", "Civilian and responder", "Incoming/active incident call signaling"),
            ("Responder presence", "Responder and dispatcher overview", "RTDB online nodes with onDisconnect; Firestore fallback"),
            ("Team on duty", "Dispatcher pages", "Command-center document subscription"),
        ], widths=[1.65, 2.3, 3.05], font_size=8.2)
        add_callout(doc, "Query resilience", "Emergency subscriptions detect missing Firestore indexes, log a single operational warning, fall back to less constrained queries, and sort/filter in memory. Authentication restoration is awaited before attaching protected listeners.")

        # PART V
        add_section_title(doc, "Part V", "Integrations and APIs", "External services, internal route handlers, environment boundaries, and communication flows.")
        doc.add_heading("19. External Integration Matrix", level=1)
        add_table(doc, ["Integration", "Purpose", "Applications", "Configuration"], [
            ("Firebase", "Identity, documents, media, presence and admin operations", "All", "NEXT_PUBLIC_/EXPO_PUBLIC_/FIREBASE variables; service account for admin routes"),
            ("Google Maps", "Native maps and markers", "Civilian and responder mobile", "GOOGLE_MAPS_API_KEY or EXPO_PUBLIC_GOOGLE_MAPS_API_KEY injected by app.config.js"),
            ("Mapbox + Leaflet", "Web tile styling and map rendering", "Dispatcher web", "NEXT_PUBLIC_MAPBOX_ACCESS_TOKEN and optional style"),
            ("Agora", "Real-time incident voice communication", "Civilian, responder, dispatcher", "AGORA_APP_ID, AGORA_APP_CERTIFICATE, token TTL; public app ID for clients"),
            ("Resend", "Verification and recovery email delivery", "Super-admin API", "RESEND_API_KEY, RESEND_FROM_EMAIL, cooldown"),
            ("Vercel", "Web build and hosting", "Dispatcher and super-admin", "Per-app vercel.json with workspace-aware install/build commands"),
            ("Expo/EAS", "Native builds and distribution", "Mobile apps", "Expo project IDs; responder development/preview/production profiles"),
        ], widths=[1.1, 2.2, 1.45, 2.25], font_size=7.9)

        doc.add_heading("20. Internal API Route Catalog", level=1)
        add_table(doc, ["Application", "Route", "Responsibility"], [
            ("Dispatcher", "/api/agora/token", "Create scoped Agora RTC tokens for incident channels"),
            ("Dispatcher", "/api/create-team-member", "Provision or associate operational team members"),
            ("Dispatcher", "/api/agent/chat", "Command-center assistant/chat request adapter"),
            ("Super Admin", "/api/create-{dispatcher,responder,civilian,command-center}", "Privileged account provisioning"),
            ("Super Admin", "/api/email-otp/{send,verify}", "Email ownership verification"),
            ("Super Admin", "/api/auth/forgot-password/{send,reset}", "Credential recovery"),
            ("Super Admin", "/api/kyc/{approve,reject}", "Civilian identity-review decisions"),
        ], widths=[1.2, 2.75, 3.05], font_size=8.1)
        doc.add_heading("20.1 API Boundary Rules", level=2)
        add_bullets(doc, [
            "Firebase Admin credentials and Resend secrets are server-side only and must never use NEXT_PUBLIC_ or EXPO_PUBLIC_ prefixes.",
            "Client applications authenticate users with Firebase Auth and rely on security rules for direct Firestore/Storage access.",
            "Privileged user creation, KYC and OTP document access execute in protected Next.js API routes.",
            "Agora clients receive short-lived tokens from the dispatcher token route; the App Certificate remains server-only.",
            "Public Firebase and map identifiers are build-time/runtime configuration values, not authorization controls; authorization remains rules- and identity-based.",
        ])

        doc.add_heading("21. Voice Calling and Messaging", level=1)
        doc.add_heading("21.1 Incident Voice Calls", level=2)
        add_numbered(doc, [
            "A caller creates a callSessions record containing incident, channel, role and target responder metadata.",
            "Subscribed responder clients receive ringing sessions and accept or decline.",
            "The application requests an Agora token for the normalized incident channel.",
            "Clients join Agora media while Firestore tracks accepted, connected, ended, missed or failed state.",
            "Cleanup functions close signaling state and release client media resources.",
        ])
        doc.add_heading("21.2 Operational Messaging", level=2)
        add_bullets(doc, [
            "Direct and group chat threads store participant identity, role and thread type.",
            "Messages are stored in chatThreads/{threadId}/messages and subscribed to in real time.",
            "Security rules permit participant reads and participant-authored creates; message mutation/deletion is denied.",
        ])

        # PART VI
        add_section_title(doc, "Part VI", "Security, Identity, and Privacy", "Authentication, authorization, media protection, privileged operations, and secure configuration.")
        doc.add_heading("22. Identity Model", level=1)
        add_table(doc, ["Identity", "Profile collection", "Authorization signal"], [
            ("Civilian", "users/{uid}", "Authenticated owner plus account status; KYC metadata on profile"),
            ("Dispatcher / responder", "dispatchers/{uid}", "Agency role, designation, active flag, and ownership"),
            ("Command center", "commandCenters/{uid}", "Command-center profile and authenticated ownership"),
            ("Super admin", "admins/{uid}", "Admin profile verified by web context/server functions"),
        ], widths=[1.6, 2.3, 3.1])
        doc.add_heading("22.1 Agency and Operational Roles", level=2)
        doc.add_paragraph("DispatcherRole supports BFP, PNP, MDRRMO, AMBULANCE and PCG. Incident agency routing expands the catalog to include RESCUE_1111, TCPGH, CHO, TFLC, PSSO_TCTMG, barangay officials, utility providers, command center and other agencies.")

        doc.add_heading("23. Firestore Authorization", level=1)
        add_table(doc, ["Data area", "Read posture", "Write posture"], [
            ("Profiles", "Owner plus authorized operational/admin roles", "Owner and/or administrative role depending on collection"),
            ("Emergency reports", "Authenticated users", "Owner/operational roles; deletion restricted to command center"),
            ("Incidents", "Authenticated users", "Operational/admin roles with explicit rule conditions"),
            ("Resources / dispatches", "Authenticated users", "Command-center/dispatcher roles; super admin included for dispatch records"),
            ("Teams", "Authenticated users", "Command-center or super admin"),
            ("Calls", "Authenticated participants/readers", "Authenticated updates; restricted create ownership and administrative delete"),
            ("Chat", "Thread/message participants", "Participant creates; message update/delete denied"),
            ("OTP collections", "No client access", "Server/Admin SDK only"),
            ("Fallback", "Denied", "Denied"),
        ], widths=[1.6, 2.5, 2.9], font_size=8.2)
        add_callout(doc, "Security posture", "The rule sets use explicit collection matches and end with deny-by-default fallbacks. Administrative API routes bypass client rules through the Firebase Admin SDK and therefore require strict server credential and endpoint protection.", "warning")

        doc.add_heading("24. Storage Authorization", level=1)
        add_table(doc, ["Path", "Write", "Read"], [
            ("kyc-documents/{uid}/{fileName}", "Authenticated owner with size/type constraints", "Owner or super admin"),
            ("emergencies/photos/{photoId}", "Authenticated authorized upload with validation", "Authenticated users"),
            ("post-reports/{incidentId}/{photoId}", "Authenticated response workflow", "Public in current rule set"),
            ("All other paths", "Denied", "Denied"),
        ], widths=[2.35, 2.35, 2.3])
        add_callout(doc, "Privacy review item", "The current post-report photo rule permits public reads. Confirm that this matches operational and privacy requirements before production deployment; otherwise restrict reads to authenticated participants or authorized operational roles.", "risk")

        doc.add_heading("25. Secret and Configuration Handling", level=1)
        add_bullets(doc, [
            "Never commit .env, .env.local, service-account JSON, API certificates or email-provider secrets.",
            "Use EXPO_PUBLIC_ only for values required in mobile client bundles; use NEXT_PUBLIC_ only for values required in browser bundles.",
            "Keep FIREBASE_SERVICE_ACCOUNT_JSON, GOOGLE_APPLICATION_CREDENTIALS, AGORA_APP_CERTIFICATE and RESEND_API_KEY server-only.",
            "Treat Firebase web configuration and map keys as identifiers protected through provider restrictions and security rules, not as secrets.",
            "Restrict Google Maps and Mapbox keys by package/bundle/domain and permitted APIs.",
            "Rotate credentials after exposure and redeploy affected clients/services.",
        ])

        # PART VII
        add_section_title(doc, "Part VII", "Configuration and Deployment", "Environment matrix, build dependencies, web deployment, mobile build profiles, and release sequence.")
        doc.add_heading("26. Environment Variable Matrix", level=1)
        env_rows = [
            ("Firebase client", "NEXT_PUBLIC_FIREBASE_API_KEY / EXPO_PUBLIC_FIREBASE_API_KEY", "Web browser / mobile bundle", "Required"),
            ("Firebase client", "*_FIREBASE_AUTH_DOMAIN, PROJECT_ID, STORAGE_BUCKET, MESSAGING_SENDER_ID, APP_ID", "Web/mobile", "Required"),
            ("Firebase presence", "NEXT_PUBLIC_FIREBASE_DATABASE_URL / EXPO_PUBLIC_FIREBASE_DATABASE_URL", "Web/mobile", "Recommended; enables RTDB presence"),
            ("Firebase scripts", "FIREBASE_API_KEY, AUTH_DOMAIN, PROJECT_ID, STORAGE_BUCKET, MESSAGING_SENDER_ID, APP_ID, DATABASE_URL", "packages/firebase scripts", "As needed"),
            ("Admin SDK", "FIREBASE_SERVICE_ACCOUNT_JSON or GOOGLE_APPLICATION_CREDENTIALS", "Super-admin server", "Required for privileged APIs"),
            ("Native maps", "GOOGLE_MAPS_API_KEY or EXPO_PUBLIC_GOOGLE_MAPS_API_KEY", "Mobile build", "Required for Google provider in native Android"),
            ("Web maps", "NEXT_PUBLIC_MAPBOX_ACCESS_TOKEN; NEXT_PUBLIC_MAPBOX_STYLE", "Dispatcher browser", "Required/optional style"),
            ("Agora server", "AGORA_APP_ID; AGORA_APP_CERTIFICATE; AGORA_TOKEN_TTL_SECONDS", "Dispatcher server", "Required for production voice"),
            ("Agora client", "EXPO_PUBLIC_AGORA_APP_ID", "Mobile bundle", "Required for voice"),
            ("Email", "RESEND_API_KEY; RESEND_FROM_EMAIL; RESEND_COOLDOWN_MS", "Super-admin server", "Required for OTP email"),
            ("API addressing", "EXPO_PUBLIC_API_URL / APP_URL / BASE_URL", "Mobile bundle", "Environment-specific"),
        ]
        add_table(doc, ["Area", "Variables", "Runtime", "Requirement"], env_rows, widths=[1.2, 3.0, 1.6, 1.2], font_size=7.45)

        doc.add_heading("27. Build and Dependency Order", level=1)
        add_numbered(doc, [
            "Install Node.js 18 or later, npm, Git, and the platform-specific mobile toolchain when native builds are required.",
            "From the repository root, run npm install so workspace dependencies and the root lockfile remain consistent.",
            "Build @packages/firebase before production web builds or after changing shared TypeScript modules.",
            "Configure application-specific environment files or deployment secrets without committing them.",
            "Run web builds and mobile configuration/build checks for the target environment.",
            "Deploy Firebase rules/indexes before releasing client features that depend on new collections or queries.",
        ])
        add_code(doc, """# Shared package
npm run build --workspace @packages/firebase

# Dispatcher web
npm run build --workspace dispatcher-web-app

# Super-admin web
npm run build --workspace super-admin-web-app

# Mobile development
npm run start --workspace civilian-mobile-app
npm exec --workspace responder-mobile-app -- expo start""")

        doc.add_heading("28. Web Deployment", level=1)
        add_bullets(doc, [
            "Dispatcher and super-admin apps each provide an app-local vercel.json.",
            "Install commands include the application workspace, @packages/firebase and the workspace root while ignoring install scripts in the remote build.",
            "Build commands compile the shared Firebase package first, then run the application build.",
            "Ignore commands redeploy only when the selected app, shared Firebase package, root package.json or package-lock.json changes.",
            "Environment variables are configured independently per Vercel project; server-only variables must not be exposed as public variables.",
        ])

        doc.add_heading("29. Mobile Build and Release", level=1)
        add_table(doc, ["Application", "Identity", "Release configuration"], [
            ("Civilian", "com.tuguegarao.resqlink; scheme resqlink", "Expo project metadata, new architecture enabled, microphone/location permissions; environment-driven maps/Firebase/Agora"),
            ("Responder", "com.tuguegarao.resqlink.responder; scheme resqlink-responder", "EAS development, preview and production profiles; production auto-increments app version"),
        ], widths=[1.2, 2.5, 3.3])
        add_bullets(doc, [
            "Validate Android and iOS permissions and user-facing descriptions before store submission.",
            "Use build-time secrets/environment profiles for Firebase, maps and Agora configuration.",
            "Verify native Google Maps behavior on a development/preview build; Expo Go behavior does not fully represent standalone native configuration.",
            "Exercise voice, location, notifications, background/foreground presence and deep-link routes on physical devices.",
        ])

        doc.add_heading("30. Firebase Deployment Assets", level=1)
        add_table(doc, ["Asset", "Purpose"], [
            ("packages/firebase/firestore.rules", "Firestore authorization policy"),
            ("packages/firebase/firestore.indexes.json", "Composite/index query support"),
            ("packages/firebase/storage.rules", "KYC, emergency and post-report media access"),
            ("packages/firebase/database.rules.json", "Realtime Database presence access"),
            ("packages/firebase/scripts/*", "Admin bootstrap, user creation, seeding, migration and lifecycle validation"),
        ], widths=[3.0, 4.0])

        # PART VIII
        add_section_title(doc, "Part VIII", "Engineering, Quality, and Operations", "Development standards, verification strategy, observability, recovery, and maintenance priorities.")
        doc.add_heading("31. Development Workflow", level=1)
        add_numbered(doc, [
            "Create a short-lived feature or integration branch from the current main baseline.",
            "Keep shared lifecycle and persistence logic in @packages/firebase; keep role-specific orchestration and presentation in the consuming app.",
            "Update security rules and indexes in the same change when introducing a new collection, access path or compound query.",
            "Run shared TypeScript compilation, affected web builds, mobile syntax/config checks and targeted behavior tests.",
            "Test role boundaries with representative civilian, dispatcher, responder, command-center and admin accounts.",
            "Document environment or deployment changes and verify rollback points before release.",
        ])
        doc.add_heading("31.1 Architectural Guardrails", level=2)
        add_bullets(doc, [
            "Do not write incident/report status fields directly when a shared domain transition function exists.",
            "Do not initialize parallel Firebase app instances; use lazy accessors from config.ts.",
            "Do not expose Admin SDK or provider secrets to browser/mobile bundles.",
            "Keep route files thin and place feature logic in screens/hooks/modules.",
            "Preserve permanent team assignment snapshots and history during reassignment.",
            "Unsubscribe Firestore/RTDB listeners and Agora sessions during unmount, sign-out and app-background transitions.",
        ])

        doc.add_heading("32. Verification Strategy", level=1)
        add_table(doc, ["Layer", "Minimum checks", "Release evidence"], [
            ("Shared package", "tsc build; lifecycle and migration scripts; rule/index review", "Successful build and script output"),
            ("Web apps", "Next production build, authenticated route smoke tests, API route tests", "Build logs and environment checklist"),
            ("Mobile apps", "Expo Doctor/config, native preview build, physical-device smoke tests", "EAS build and device test record"),
            ("Cross-role", "Report → triage → dispatch → response → closure", "Scenario checklist with record IDs/timestamps"),
            ("Security", "Positive and negative rule tests by role and ownership", "Emulator/rules test results"),
            ("Resilience", "Offline/reconnect, missing index, missing map key, denied location, call failure", "Observed fallback behavior"),
            ("Performance", "Large active lists, map marker volume, listener cleanup, export size", "Timing and memory observations"),
        ], widths=[1.2, 3.65, 2.15], font_size=8.1)

        doc.add_heading("33. Operational Observability", level=1)
        add_bullets(doc, [
            "Client and shared modules currently use structured console logging helpers, including once-only warnings for repeated Firebase initialization/index conditions.",
            "Critical operational data is observable through record timestamps, status fields, acknowledgment metadata, response times, team assignment history and call-session state.",
            "Vercel provides server/API build and runtime logs; EAS provides mobile build logs; Firebase consoles provide Auth, Firestore, Storage and RTDB inspection.",
            "Production hardening should add centralized error reporting, release/version tags, API correlation IDs, audit-event retention and alert thresholds.",
        ])

        doc.add_heading("34. Troubleshooting Runbook", level=1)
        add_table(doc, ["Symptom", "Likely cause", "First actions"], [
            ("Cannot resolve @packages/firebase", "Shared package not installed/built or workspace resolution issue", "Run root npm install; build @packages/firebase; verify Metro/Next workspace settings"),
            ("Firebase config incomplete", "Missing/mis-prefixed environment values", "Check app-specific env file/deployment settings and prefixes; restart/rebuild"),
            ("Permission denied", "Rule mismatch, wrong profile role/status or unauthenticated request", "Inspect UID/profile collection, auth token and relevant Firestore/Storage rule"),
            ("Missing-index warning", "Compound query index not deployed", "Deploy firestore.indexes.json; temporary in-memory fallback may remain operational"),
            ("Map unavailable on Android", "Google Maps key not injected/restricted incorrectly", "Set build variable, verify app.config output and package restriction, rebuild native app"),
            ("Responder count remains zero", "RTDB URL/rules missing or designation mismatch", "Check database URL, presence rules, profile designation and Firestore fallback"),
            ("Voice call fails", "Agora credentials/token/channel mismatch or microphone permission", "Inspect token route, app ID, certificate, TTL, channel normalization and device permission"),
            ("OTP email not delivered", "Resend configuration, cooldown or sender verification", "Check server logs, API key, from address, recipient and OTP document"),
            ("Web build fails", "Shared package, environment or workspace dependency mismatch", "Build Firebase first, clean .next, reinstall from lockfile and run app build"),
        ], widths=[1.65, 2.45, 2.9], font_size=7.9)

        doc.add_heading("35. Backup, Recovery, and Continuity", level=1)
        add_bullets(doc, [
            "Enable scheduled Firestore exports and define retention/recovery ownership for production projects.",
            "Protect service-account credentials and maintain break-glass access through audited secret management.",
            "Tag production releases and retain the web/mobile configuration used for each release.",
            "Document how to disable compromised accounts, rotate provider credentials and revoke active sessions.",
            "Exercise recovery of Firestore data, Storage objects, Vercel deployments and mobile release configurations.",
            "Use incident and team assignment history to reconstruct operational actions after service restoration.",
        ])

        doc.add_heading("36. Current-State Considerations and Recommendations", level=1)
        recommendation_rows = [
            ("Documentation drift", "Some legacy READMEs describe earlier versions or mock-only states", "Treat current code and this baseline as authoritative; revise per-app READMEs during future releases", "High"),
            ("Automated testing", "Jest dependencies exist, but end-to-end and rule-test coverage is not demonstrated across all apps", "Add emulator-backed domain/rules tests and cross-role E2E scenarios", "High"),
            ("Observability", "Console-based logging dominates", "Integrate centralized error/performance monitoring and correlation IDs", "High"),
            ("Storage privacy", "Post-report photos are publicly readable in current storage rules", "Confirm requirement and restrict if not explicitly public", "High"),
            ("Dependency health", "Workspace uses extensive overrides and patching; lint tooling can be sensitive to install state", "Use clean reproducible CI installs, dependency audits and scheduled upgrades", "Medium"),
            ("Mobile parity", "Responder has explicit EAS profiles; civilian relies on project metadata/config", "Define equivalent checked-in release profiles and environment policy", "Medium"),
            ("Auditability", "Operational timestamps exist but a unified audit-event stream is not present", "Add immutable administrative/operational audit events", "Medium"),
            ("Performance", "Real-time maps and active lists may grow", "Adopt pagination/windowing, listener scoping and production volume testing", "Medium"),
        ]
        add_table(doc, ["Area", "Observation", "Recommendation", "Priority"], recommendation_rows, widths=[1.25, 2.15, 3.0, 0.6], font_size=7.45)

        # PART IX
        add_section_title(doc, "Part IX", "Reference Appendices", "Routes, configuration quick reference, terminology, ownership, and repository evidence.")
        doc.add_heading("Appendix A. Route Inventory", level=1)
        add_table(doc, ["Application", "Route families"], [
            ("Civilian mobile", "(auth): login, register, email verification, account pending, forgot/reset password; (main): dashboard, history, profile, emergency form/confirmation, responder map, calling; (settings): appearance, notifications, privacy/security, FAQ, support, report issue"),
            ("Responder mobile", "auth/login; tabs/dashboard, map, notifications, settings; incident/[id]; support/about, help-support, location"),
            ("Dispatcher web", "overview, intake, incidents, incident-management, map, resources, teams, history, report, report/incidents, footage-requests, login"),
            ("Super-admin web", "dashboard, responders, dispatchers, civilians, command-centers, KYC, login"),
        ], widths=[1.5, 5.5], font_size=8.3)

        doc.add_heading("Appendix B. Operational Status Quick Reference", level=1)
        add_table(doc, ["Phase", "Emergency report", "Incident", "Resource", "Responder action"], [
            ("Intake", "pending", "new", "available", "No assignment"),
            ("Needs capacity", "pending/linked", "awaiting_resources or liaison_pending", "available", "Await assignment"),
            ("Dispatched", "active/linked", "dispatched", "assigned", "Accept or decline"),
            ("Travel", "enroute", "enroute", "en_route", "Navigate to scene"),
            ("Scene", "on_scene", "on_scene", "on_scene", "Touchdown and assessment"),
            ("Closure", "done/resolved", "resolved/unresolved", "available", "Submit post-report and close"),
        ], widths=[1.0, 1.45, 1.8, 1.25, 1.5], font_size=8)

        doc.add_heading("Appendix C. Glossary", level=1)
        glossary = [
            ("BFP", "Bureau of Fire Protection"),
            ("PNP", "Philippine National Police"),
            ("MDRRMO", "Municipal Disaster Risk Reduction and Management Office"),
            ("PCG", "Philippine Coast Guard"),
            ("KYC", "Know Your Customer / identity verification review"),
            ("Emergency report", "Original reporter-facing record, usually in emergencies"),
            ("Incident", "Canonical command-center operational record in incidents"),
            ("Dispatch record", "Incident-to-agency/resource relationship in incidentDispatches"),
            ("Operational team", "Team assigned permanent ownership for an incident"),
            ("Touchdown", "Recorded responder arrival/on-scene event, optionally with distance/source"),
            ("RTDB", "Firebase Realtime Database, used primarily for responder presence"),
            ("EAS", "Expo Application Services for mobile builds and submission"),
            ("OTP", "One-time passcode used for verification or password recovery"),
        ]
        add_table(doc, ["Term", "Definition"], glossary, widths=[1.6, 5.4])

        doc.add_heading("Appendix D. Repository Evidence", level=1)
        source_rows = [
            ("Workspace configuration", "package.json; package-lock.json"),
            ("Civilian architecture", "apps/civilian-mobile-app/src; app.json; app.config.js; docs/*"),
            ("Responder architecture", "apps/responder-mobile-app/src; app.json; app.config.js; eas.json; docs/*"),
            ("Dispatcher architecture", "apps/dispatcher-web-app/app; components; contexts; lib; vercel.json; docs/*"),
            ("Super-admin architecture", "apps/super-admin-web-app/app; contexts; lib; vercel.json; README.md"),
            ("Domain/data layer", "packages/firebase/src/*; package.json"),
            ("Security", "packages/firebase/firestore.rules; storage.rules; database.rules.json"),
            ("Deployment", "apps/*/vercel.json; apps/responder-mobile-app/eas.json; Expo app configuration"),
            ("Branding", "apps/dispatcher-web-app/public/branding/resq-link-logo.png and resq-link-icon.png"),
            ("Live screenshots", "docs/screenshots/dispatcher-web-login-live.png; super-admin-live.png; civilian-mobile-live.png; responder-mobile-live.png"),
        ]
        add_table(doc, ["Evidence area", "Repository source"], source_rows, widths=[2.0, 5.0], font_size=8.3)

        doc.add_heading("Appendix E. Documentation Maintenance Checklist", level=1)
        add_bullets(doc, [
            "Update the repository baseline, version and revision history for every approved documentation release.",
            "Reconcile route inventories after adding, moving or deleting screens/pages/API handlers.",
            "Update collection, status and security sections alongside domain/rule changes.",
            "Update environment and deployment sections whenever a provider, variable, project ID or build profile changes.",
            "Record new operational risks, mitigations and runbook steps after production incidents or exercises.",
            "Regenerate the table of contents and verify page numbering, diagrams, logo quality and table wrapping before distribution.",
        ])

        end = doc.add_paragraph()
        end.alignment = WD_ALIGN_PARAGRAPH.CENTER
        end.paragraph_format.space_before = Pt(35)
        if ICON.exists():
            end.add_run().add_picture(str(ICON), width=Inches(0.55))
        end2 = doc.add_paragraph()
        end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = end2.add_run("RESQ-Link")
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = rgb(PRIMARY)
        end3 = doc.add_paragraph()
        end3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        end3.add_run("Connected response. Shared awareness. Accountable outcomes.").italic = True
        end3.runs[0].font.color.rgb = rgb(MUTED)

    return doc


if __name__ == "__main__":
    if not LOGO.exists():
        raise FileNotFoundError(f"Brand logo not found: {LOGO}")
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)
